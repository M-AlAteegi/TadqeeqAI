"""
TadqeeqAI v3.0 - Backend
========================
All business logic, classes, and API endpoints.

v3.0 Enhancements:
- Lazy imports for heavy libraries
- Background threading for initialization
- Real-time progress reporting (stage + percentage)
- GPU detection with OLLAMA_VULKAN auto-configuration
- Ollama availability checking with user-friendly prompts
"""

import json
import pickle
import re
import os
import subprocess
import time
import uuid
import warnings
import logging
import atexit
import base64
import threading
from datetime import datetime
from pathlib import Path
from io import BytesIO
import numpy as np

warnings.filterwarnings('ignore', category=DeprecationWarning)
logging.getLogger('pywebview').setLevel(logging.ERROR)

# ==============================================================================
# CONFIGURATION
# ==============================================================================

CHROMA_PATH = "./chroma_db_v2"
BM25_PATH = "./bm25_index.pkl"
DOCS_PATH = "./documents.json"
CHAT_HISTORY_PATH = "./chat_history"
EMBEDDING_MODEL = 'intfloat/multilingual-e5-base'
MAX_DOCUMENT_PAGES = 50
LLM_MODEL = 'aya:8b'
OLLAMA_DOWNLOAD_URL = 'https://ollama.com/download'

Path(CHAT_HISTORY_PATH).mkdir(exist_ok=True)
OLLAMA_PROCESS = None

# ==============================================================================
# PROGRESS TRACKING
# ==============================================================================

class ProgressTracker:
    STAGES = [
        ('gpu_detect', 'Detecting GPU...', 10),
        ('ollama_check', 'Checking Ollama...', 20),
        ('ollama_start', 'Starting Ollama...', 35),
        ('documents', 'Loading documents...', 50),
        ('bm25', 'Loading BM25 index...', 60),
        ('embeddings', 'Loading embedding model...', 80),
        ('chromadb', 'Connecting to ChromaDB...', 95),
        ('ready', 'Ready!', 100),
    ]
    
    def __init__(self):
        self._lock = threading.Lock()
        self._stage = 'starting'
        self._stage_text = 'Initializing...'
        self._progress = 0
        self._error = None
        self._details = None
        self.last_brief = None
    
    def set_stage(self, stage_id, details=None):
        with self._lock:
            for sid, text, progress in self.STAGES:
                if sid == stage_id:
                    self._stage = sid
                    self._stage_text = text
                    self._progress = progress
                    self._details = details
                    print(f"  [{progress:3d}%] {text}" + (f" ({details})" if details else ""))
                    break
    
    def set_error(self, error_msg):
        with self._lock:
            self._error = error_msg
            print(f"  [ERROR] {error_msg}")
    
    def get_status(self):
        with self._lock:
            return {
                'stage': self._stage,
                'stage_text': self._stage_text,
                'progress': self._progress,
                'error': self._error,
                'details': self._details
            }

progress_tracker = ProgressTracker()

# ==============================================================================
# GPU DETECTION
# ==============================================================================

def detect_gpu_type():
    if os.name != 'nt':
        return 'unknown'
    try:
        output = subprocess.check_output(
            ["wmic", "path", "win32_VideoController", "get", "name"],
            universal_newlines=True,
            creationflags=subprocess.CREATE_NO_WINDOW
        ).lower()
        igpu_markers = ['intel', 'amd radeon graphics', 'vega', 'radeon(tm) graphics']
        dgpu_markers = ['nvidia', 'rtx', 'gtx', 'geforce', 'radeon rx', 'radeon pro']
        has_igpu = any(m in output for m in igpu_markers)
        has_dgpu = any(m in output for m in dgpu_markers)
        if has_dgpu:
            return 'dgpu'
        elif has_igpu:
            return 'igpu'
        return 'unknown'
    except Exception as e:
        print(f"    GPU detection failed: {e}")
        return 'unknown'

def configure_gpu_environment():
    gpu_type = detect_gpu_type()
    if gpu_type == 'igpu':
        os.environ['OLLAMA_VULKAN'] = '1'
        os.environ['OLLAMA_NUM_PARALLEL'] = '1'
        os.environ['OLLAMA_MAX_LOADED_MODELS'] = '1'
        print(f"    ✓ Configured for iGPU (Vulkan enabled)")
        return 'igpu'
    elif gpu_type == 'dgpu':
        os.environ.pop('OLLAMA_VULKAN', None)
        print(f"    ✓ Configured for discrete GPU")
        return 'dgpu'
    else:
        print(f"    ⚠ GPU type unknown, using defaults")
        return 'unknown'

# ==============================================================================
# OLLAMA MANAGEMENT
# ==============================================================================

def check_ollama_installed():
    if os.name == 'nt':
        ollama_paths = [
            os.path.expandvars(r'%LOCALAPPDATA%\Programs\Ollama\ollama.exe'),
            os.path.expandvars(r'%LOCALAPPDATA%\Ollama\ollama.exe'),
            r'C:\Program Files\Ollama\ollama.exe',
        ]
        for path in ollama_paths:
            if os.path.exists(path):
                return True, path
        try:
            result = subprocess.run(['where', 'ollama'], capture_output=True, creationflags=subprocess.CREATE_NO_WINDOW)
            if result.returncode == 0:
                return True, 'ollama'
        except:
            pass
        return False, None
    else:
        try:
            result = subprocess.run(['which', 'ollama'], capture_output=True)
            if result.returncode == 0:
                return True, result.stdout.decode().strip()
        except:
            pass
        return False, None

def start_ollama():
    global OLLAMA_PROCESS
    installed, ollama_path = check_ollama_installed()
    if not installed:
        return False, 'not_installed'
    
    if os.name == 'nt':
        for proc_name in ['ollama app.exe', 'Ollama.exe', 'ollama.exe']:
            subprocess.run(['taskkill', '/F', '/IM', proc_name], capture_output=True, creationflags=subprocess.CREATE_NO_WINDOW)
        time.sleep(2)
    
    try:
        if os.name == 'nt':
            startupinfo = subprocess.STARTUPINFO()
            startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW | subprocess.STARTF_USESTDHANDLES
            startupinfo.wShowWindow = 0
            OLLAMA_PROCESS = subprocess.Popen(
                [ollama_path, 'serve'],
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, stdin=subprocess.DEVNULL,
                startupinfo=startupinfo, env=os.environ,
                creationflags=subprocess.CREATE_NO_WINDOW | subprocess.DETACHED_PROCESS | subprocess.CREATE_NEW_PROCESS_GROUP
            )
        else:
            OLLAMA_PROCESS = subprocess.Popen(['ollama', 'serve'], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, env=os.environ)
        
        atexit.register(stop_ollama)
        import socket
        for i in range(30):
            time.sleep(1)
            try:
                sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                result = sock.connect_ex(('127.0.0.1', 11434))
                sock.close()
                if result == 0:
                    return True, 'started'
            except:
                pass
        return False, 'timeout'
    except Exception as e:
        return False, str(e)

def stop_ollama():
    global OLLAMA_PROCESS
    if OLLAMA_PROCESS:
        try:
            OLLAMA_PROCESS.terminate()
            OLLAMA_PROCESS.wait(timeout=5)
        except:
            try:
                OLLAMA_PROCESS.kill()
            except:
                pass
        OLLAMA_PROCESS = None
    if os.name == 'nt':
        subprocess.run(['taskkill', '/F', '/IM', 'ollama.exe'], capture_output=True, creationflags=subprocess.CREATE_NO_WINDOW)

def pull_model_with_progress(model_name, progress_callback=None):
    try:
        import ollama
        current_status = ""
        for progress in ollama.pull(model_name, stream=True):
            status = progress.get('status', '')
            if status != current_status:
                current_status = status
                if progress_callback:
                    completed = progress.get('completed', 0)
                    total = progress.get('total', 0)
                    pct = int((completed / total) * 100) if total > 0 else 0
                    progress_callback(status, pct)
        return True, None
    except Exception as e:
        return False, str(e)

# ==============================================================================
# OPTIONAL IMPORTS
# ==============================================================================

HAS_PYMUPDF = None
HAS_DOCX = None
HAS_REPORTLAB = None

def check_optional_imports():
    global HAS_PYMUPDF, HAS_DOCX, HAS_REPORTLAB
    try:
        import fitz
        HAS_PYMUPDF = True
    except ImportError:
        HAS_PYMUPDF = False
    try:
        from docx import Document
        HAS_DOCX = True
    except ImportError:
        HAS_DOCX = False
    try:
        from reportlab.lib import colors
        HAS_REPORTLAB = True
    except ImportError:
        HAS_REPORTLAB = False

# ==============================================================================
# CHAT HISTORY
# ==============================================================================

class ChatHistory:
    def __init__(self):
        self.history_path = Path(CHAT_HISTORY_PATH)
        self.current_chat_id = None
        self.current_messages = []
    
    def get_messages(self):
        return self.current_messages
    
    def new_chat(self):
        self.current_chat_id = str(uuid.uuid4())[:8]
        self.current_messages = []
        return self.current_chat_id
    
    def add_message(self, role, content, sources=None, regulator=None):
        msg = {'role': role, 'content': content, 'timestamp': datetime.now().isoformat()}
        if sources:
            msg['sources'] = sources
        if regulator:
            msg['regulator'] = regulator
        self.current_messages.append(msg)
        self._save_current()
    
    def _save_current(self):
        if not self.current_chat_id:
            return
        filepath = self.history_path / f"{self.current_chat_id}.json"
        chat_data = {
            'id': self.current_chat_id,
            'created': self.current_messages[0]['timestamp'] if self.current_messages else datetime.now().isoformat(),
            'updated': datetime.now().isoformat(),
            'messages': self.current_messages,
            'preview': self._get_preview()
        }
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(chat_data, f, ensure_ascii=False, indent=2)
    
    def _get_preview(self):
        for msg in self.current_messages:
            if msg['role'] == 'user':
                text = msg['content']
                return text[:50] + '...' if len(text) > 50 else text
        return 'New Chat'
    
    def get_recent_chats(self, limit=20):
        chats = []
        for filepath in sorted(self.history_path.glob('*.json'), key=lambda x: x.stat().st_mtime, reverse=True):
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    chats.append({'id': data['id'], 'preview': data.get('preview', 'Chat'), 'updated': data.get('updated', ''), 'message_count': len(data.get('messages', []))})
            except:
                pass
            if len(chats) >= limit:
                break
        return chats
    
    def load_chat(self, chat_id):
        filepath = self.history_path / f"{chat_id}.json"
        if filepath.exists():
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
                self.current_chat_id = chat_id
                self.current_messages = data.get('messages', [])
                return self.current_messages
        return []
    
    def get_conversation_context(self, limit=6):
        return self.current_messages[-limit:] if self.current_messages else []
    
    def delete_chat(self, chat_id):
        filepath = self.history_path / f"{chat_id}.json"
        if filepath.exists():
            filepath.unlink()
            if self.current_chat_id == chat_id:
                self.current_chat_id = None
                self.current_messages = []
            return True
        return False

# ==============================================================================
# DOCUMENT PROCESSOR
# ==============================================================================

class DocumentProcessor:
    def __init__(self):
        self.current_document = None
        self.current_text = None
        self.current_filename = None
        self.current_page_count = 0
    
    def process_file(self, file_path=None, file_data=None, filename=None):
        if HAS_PYMUPDF is None:
            check_optional_imports()
        try:
            if file_data:
                data = base64.b64decode(file_data)
                ext = filename.lower().split('.')[-1] if filename else ''
            elif file_path:
                with open(file_path, 'rb') as f:
                    data = f.read()
                ext = file_path.lower().split('.')[-1]
                filename = os.path.basename(file_path)
            else:
                return {'error': 'No file provided'}
            if ext == 'pdf':
                return self._process_pdf(data, filename)
            elif ext in ['docx', 'doc']:
                return self._process_docx(data, filename)
            else:
                return {'error': f'Unsupported file type: {ext}'}
        except Exception as e:
            return {'error': str(e)}
    
    def _process_pdf(self, data, filename):
        if not HAS_PYMUPDF:
            return {'error': 'PDF support not available. Install PyMuPDF: pip install PyMuPDF'}
        import fitz
        try:
            doc = fitz.open(stream=data, filetype="pdf")
            if doc.page_count > MAX_DOCUMENT_PAGES:
                return {'error': f'Document too large. Maximum {MAX_DOCUMENT_PAGES} pages allowed.'}
            text_parts = []
            for page_num, page in enumerate(doc):
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")
            full_text = "\n\n".join(text_parts)
            self.current_document = doc
            self.current_text = full_text
            self.current_filename = filename
            self.current_page_count = doc.page_count
            summary = self._generate_quick_summary(full_text, filename)
            return {'success': True, 'filename': filename, 'pages': doc.page_count, 'chars': len(full_text), 'summary': summary}
        except Exception as e:
            return {'error': f'Failed to process PDF: {str(e)}'}
    
    def _process_docx(self, data, filename):
        if not HAS_DOCX:
            return {'error': 'DOCX support not available. Install python-docx: pip install python-docx'}
        from docx import Document as DocxDocument
        try:
            doc = DocxDocument(BytesIO(data))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            full_text = "\n\n".join(text_parts)
            self.current_text = full_text
            self.current_filename = filename
            self.current_page_count = len(text_parts) // 20 + 1
            summary = self._generate_quick_summary(full_text, filename)
            return {'success': True, 'filename': filename, 'pages': self.current_page_count, 'chars': len(full_text), 'summary': summary}
        except Exception as e:
            return {'error': f'Failed to process DOCX: {str(e)}'}
    
    def _generate_quick_summary(self, text, filename):
        text_lower = text.lower()
        doc_type = "Document"
        if any(kw in text_lower for kw in ['fund', 'investment', 'investor', 'subscription']):
            doc_type = "Investment Fund Document"
        elif any(kw in text_lower for kw in ['sukuk', 'bond', 'debt instrument']):
            doc_type = "Sukuk/Debt Instrument Document"
        elif any(kw in text_lower for kw in ['license', 'licensing', 'authorization']):
            doc_type = "Licensing Document"
        elif any(kw in text_lower for kw in ['contract', 'agreement', 'party', 'parties']):
            doc_type = "Contract/Agreement"
        elif any(kw in text_lower for kw in ['prospectus', 'offering', 'securities']):
            doc_type = "Securities Prospectus"
        has_arabic = bool(re.search(r'[\u0600-\u06FF]', text))
        return {'type': doc_type, 'has_arabic': has_arabic, 'word_count': len(text.split())}
    
    def get_current_text(self):
        return self.current_text
    
    def clear(self):
        self.current_document = None
        self.current_text = None
        self.current_filename = None
        self.current_page_count = 0

# ==============================================================================
# COMPLIANCE CHECKER
# ==============================================================================

class ComplianceChecker:
    COMPLIANCE_CATEGORIES = {
        'qualified_investor': {'name': 'Qualified Investor Definition', 'keywords': ['qualified investor', 'accredited investor', 'مستثمر مؤهل'], 'regulation': 'CMA Rules on Offer of Securities, Article 15', 'description': 'Documents offering securities must define qualified investor criteria'},
        'risk_disclosure': {'name': 'Risk Disclosure', 'keywords': ['risk', 'risks', 'risk factors', 'مخاطر'], 'regulation': 'CMA Rules on Offer of Securities, Article 22', 'description': 'Offering documents must include comprehensive risk disclosures'},
        'capital_requirements': {'name': 'Capital Requirements', 'keywords': ['minimum capital', 'paid-up capital', 'رأس المال'], 'regulation': 'Finance Companies Control Law, Article 5', 'description': 'Finance companies must meet minimum capital requirements'},
        'license_reference': {'name': 'Licensing Information', 'keywords': ['license', 'licensed', 'CMA', 'SAMA', 'ترخيص'], 'regulation': 'Capital Market Law, Article 3', 'description': 'Financial activities require proper licensing'},
        'fund_terms': {'name': 'Fund Terms & Conditions', 'keywords': ['terms and conditions', 'subscription', 'management fee', 'شروط وأحكام'], 'regulation': 'Investment Funds Regulations, Article 20', 'description': 'Investment funds must clearly state terms and fees'},
        'disclosure_requirements': {'name': 'Disclosure Requirements', 'keywords': ['disclosure', 'material information', 'إفصاح'], 'regulation': 'CMA Rules on Offer of Securities, Article 30', 'description': 'Issuers must disclose all material information'}
    }
    
    def check_compliance(self, text, filename=None):
        results = {'filename': filename or 'Document', 'timestamp': datetime.now().isoformat(), 'checks': [], 'summary': {'compliant': 0, 'warnings': 0, 'missing': 0}}
        text_lower = text.lower()
        for category_id, category in self.COMPLIANCE_CATEGORIES.items():
            found_keywords = [kw for kw in category['keywords'] if kw.lower() in text_lower or kw in text]
            if found_keywords:
                status = 'compliant'
                results['summary']['compliant'] += 1
                detail = f"Found references: {', '.join(found_keywords[:3])}"
            else:
                status = 'warning'
                results['summary']['warnings'] += 1
                detail = f"Consider adding {category['name'].lower()} information"
            results['checks'].append({'id': category_id, 'name': category['name'], 'status': status, 'regulation': category['regulation'], 'description': category['description'], 'detail': detail})
        total = results['summary']['compliant'] + results['summary']['warnings']
        results['score'] = round((results['summary']['compliant'] / total) * 100) if total > 0 else 100
        return results

# ==============================================================================
# CHAT EXPORTER
# ==============================================================================

class ChatExporter:
    def _sanitize_for_pdf(self, text):
        import html
        text = html.escape(text)
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
        text = re.sub(r'^[\-\*]\s+', '- ', text, flags=re.MULTILINE)
        text = text.replace('<', '(').replace('>', ')')
        return text
    
    def export_markdown(self, messages, filename=None):
        if not messages:
            return None, "No messages to export"
        md_lines = ["# TadqeeqAI Chat Export", f"**Date:** {datetime.now().strftime('%Y-%m-%d %H:%M')}", "", "---", ""]
        for msg in messages:
            role = "User" if msg['role'] == 'user' else "TadqeeqAI"
            md_lines.extend([f"## {role}", "", msg['content'], ""])
            if msg.get('sources'):
                md_lines.append("**Sources:**")
                for src in msg['sources']:
                    md_lines.append(f"- {src['article']} ({src['document']})")
                md_lines.append("")
            md_lines.extend(["---", ""])
        md_lines.extend(["", "*Generated by TadqeeqAI v3.0*"])
        return "\n".join(md_lines), None
    
    def export_pdf(self, messages, filename=None):
        if not messages:
            return None, "No messages to export"
        if HAS_REPORTLAB is None:
            check_optional_imports()
        if not HAS_REPORTLAB:
            return None, "PDF export not available. Install reportlab: pip install reportlab"
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.enums import TA_CENTER
        try:
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=24, spaceAfter=20, textColor=colors.HexColor('#00d4aa'))
            normal_style = ParagraphStyle('CustomNormal', parent=styles['Normal'], fontSize=11, leading=16)
            role_style = ParagraphStyle('CustomRole', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#00d4aa'), fontName='Helvetica-Bold', spaceBefore=15)
            footer_style = ParagraphStyle('CustomFooter', parent=styles['Normal'], fontSize=8, textColor=colors.grey, alignment=TA_CENTER)
            story = [Paragraph("TadqeeqAI Chat Export", title_style), Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}", styles['Normal']), HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e0e0e0')), Spacer(1, 20)]
            for msg in messages:
                role = "You" if msg['role'] == 'user' else "TadqeeqAI"
                story.append(Paragraph(role, role_style))
                content = self._sanitize_for_pdf(msg['content'])
                for para in content.split('\n'):
                    para = para.strip()
                    if para:
                        try:
                            story.append(Paragraph(para, normal_style))
                        except:
                            story.append(Paragraph(para.encode('ascii', 'ignore').decode(), normal_style))
                story.append(Spacer(1, 10))
                if msg.get('sources'):
                    sources_text = "Sources: " + ", ".join([s['article'] for s in msg['sources'][:3]])
                    story.append(Paragraph(sources_text, ParagraphStyle('Sources', parent=styles['Normal'], fontSize=9, textColor=colors.grey)))
                    story.append(Spacer(1, 5))
            story.extend([Spacer(1, 30), HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e0e0e0')), Spacer(1, 10), Paragraph("Generated by TadqeeqAI v3.0", footer_style)])
            doc.build(story)
            pdf_data = buffer.getvalue()
            buffer.close()
            return base64.b64encode(pdf_data).decode('utf-8'), None
        except Exception as e:
            return None, f"PDF export failed: {str(e)}"
        
    def export_brief_markdown(self, text):
        if not text: return None, "No brief available to export."
        return text, None

    def export_brief_pdf(self, text):
        if not text: return None, "No brief available to export."
        if not HAS_REPORTLAB: return None, "ReportLab not installed."
        
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from io import BytesIO
        import base64
        import re

        try:
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            styles = getSampleStyleSheet()
            # Custom Styles for the Report
            styles.add(ParagraphStyle(name='BriefTitle', parent=styles['Heading1'], fontSize=20, spaceAfter=20, textColor=colors.HexColor('#00d4aa')))
            styles.add(ParagraphStyle(name='BriefHeader', parent=styles['Heading2'], fontSize=14, spaceBefore=15, spaceAfter=10, textColor=colors.HexColor('#2c3e50')))
            
            story = [Paragraph("Executive Brief", styles['BriefTitle']), HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e0e0e0')), Spacer(1, 20)]
            
            # Simple Markdown parsing for PDF
            for line in text.split('\n'):
                line = line.strip()
                if not line: continue
                
                if line.startswith('# '): continue # Skip main title as we added it
                
                if line.startswith('## '):
                    story.append(Paragraph(line.replace('## ', '').replace('🚨', '').replace('💰', '').replace('📅', '').strip(), styles['BriefHeader']))
                elif line.startswith('* ') or line.startswith('- '):
                    # List items
                    clean_line = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', line[2:]) # Handle bold
                    story.append(Paragraph(f"• {clean_line}", styles['Normal']))
                else:
                    clean_line = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', line)
                    story.append(Paragraph(clean_line, styles['Normal']))
                
                story.append(Spacer(1, 6))

            doc.build(story)
            pdf_data = buffer.getvalue()
            buffer.close()
            return base64.b64encode(pdf_data).decode('utf-8'), None
        except Exception as e:
            return None, str(e)

# ==============================================================================
# END OF PART 1 - PART 2 CONTAINS: TadqeeqRAG class and API class
# ==============================================================================


class TadqeeqRAG:
    """Hybrid RAG system with BM25 + semantic search."""
    _instance = None
    _initialized = False
    _init_error = None
    
    @classmethod
    def get_instance(cls):
        if cls._instance is None:
            cls._instance = cls()
        return cls._instance
    
    @classmethod
    def is_ready(cls):
        return cls._instance is not None and cls._initialized
    
    def __init__(self):
        self.documents = None
        self.bm25 = None
        self.embedder = None
        self.client = None
        self.collection = None
        self.chat_history = None
        self.doc_processor = None
        self.compliance_checker = None
        self.chat_exporter = None
        self.stats = None
        self.sama_count = 0
        self.cma_count = 0
        self.total = 0
    
    def initialize(self):
        """Initialize the RAG system with progress tracking."""
        global progress_tracker
        try:
            print("\nLoading TadqeeqAI v3.0...")
            
            # Stage 1: GPU Detection
            progress_tracker.set_stage('gpu_detect')
            gpu_type = configure_gpu_environment()
            
            # Stage 2: Check Ollama
            progress_tracker.set_stage('ollama_check')
            installed, ollama_path = check_ollama_installed()
            if not installed:
                progress_tracker.set_error(f'Ollama not installed. Please download from: {OLLAMA_DOWNLOAD_URL}')
                TadqeeqRAG._init_error = {'type': 'ollama_not_installed', 'message': 'Ollama is not installed on this system.', 'download_url': OLLAMA_DOWNLOAD_URL}
                return False
            
            # Stage 3: Start Ollama
            progress_tracker.set_stage('ollama_start')
            success, status = start_ollama()
            if not success and status == 'not_installed':
                progress_tracker.set_error('Ollama not found')
                TadqeeqRAG._init_error = {'type': 'ollama_not_installed', 'message': 'Ollama is not installed.', 'download_url': OLLAMA_DOWNLOAD_URL}
                return False
            
            # Stage 4: Load Documents
            progress_tracker.set_stage('documents')
            with open(DOCS_PATH, 'r', encoding='utf-8') as f:
                self.documents = json.load(f)
            self.stats = {'SAMA': {'en': 0, 'ar': 0}, 'CMA': {'en': 0, 'ar': 0}}
            for doc in self.documents:
                reg = doc.get('regulator', 'CMA')
                lang = doc.get('language', 'en')
                self.stats[reg][lang] += 1
            self.sama_count = self.stats['SAMA']['en'] + self.stats['SAMA']['ar']
            self.cma_count = self.stats['CMA']['en'] + self.stats['CMA']['ar']
            self.total = len(self.documents)
            progress_tracker.set_stage('documents', f'{self.total} articles')
            
            # Stage 5: Load BM25
            progress_tracker.set_stage('bm25')
            with open(BM25_PATH, 'rb') as f:
                self.bm25 = pickle.load(f)
            
            # Stage 6: Load Embeddings (LAZY IMPORT)
            progress_tracker.set_stage('embeddings')
            from sentence_transformers import SentenceTransformer
            self.embedder = SentenceTransformer(EMBEDDING_MODEL)
            
            # Stage 7: Connect to ChromaDB (LAZY IMPORT)
            progress_tracker.set_stage('chromadb')
            import chromadb
            self.client = chromadb.PersistentClient(path=CHROMA_PATH)
            self.collection = self.client.get_collection("tadqeeq_v2")
            progress_tracker.set_stage('chromadb', f'{self.collection.count()} vectors')
            
            # --- REMOVED STAGE 8 (LLM WARMUP) ---
            
            # Initialize helper classes
            self.chat_history = ChatHistory()
            self.doc_processor = DocumentProcessor()
            self.compliance_checker = ComplianceChecker()
            self.chat_exporter = ChatExporter()
            
            progress_tracker.set_stage('ready')
            TadqeeqRAG._initialized = True
            print("\n✓ TadqeeqAI v3.0 Ready!\n")
            return True
        except Exception as e:
            import traceback
            traceback.print_exc()
            progress_tracker.set_error(str(e))
            TadqeeqRAG._init_error = {'type': 'init_error', 'message': str(e)}
            return False
    
    
    def detect_language(self, text):
        arabic_chars = len(re.findall(r'[\u0600-\u06FF]', text))
        return 'ar' if arabic_chars > len(text) * 0.3 else 'en'
    
    def detect_regulator(self, query):
        q_lower = query.lower()
        sama_en = ['sama', 'finance company', 'finance companies', 'financing company', 'licensing fee', 'real estate finance', 'mortgage', 'microfinance', 'finance control', 'monetary authority', 'bank', 'banking']
        cma_en = ['cma', 'capital market', 'securities', 'sukuk', 'debt instrument', 'investment fund', 'qualified investor', 'public offering', 'ipo', 'private placement', 'prospectus', 'listing', 'merger', 'acquisition', 'stock', 'shares', 'exchange']
        sama_ar = ['ساما', 'شركة التمويل', 'شركات التمويل', 'رسوم الترخيص', 'التمويل العقاري', 'التمويل الأصغر', 'مؤسسة النقد', 'البنك المركزي', 'تمويل']
        cma_ar = ['مستثمر مؤهل', 'المستثمر المؤهل', 'هيئة السوق المالية', 'هيئة السوق', 'صكوك', 'الصكوك', 'طرح عام', 'طرح خاص', 'نشرة الإصدار', 'صناديق الاستثمار', 'أوراق مالية', 'الأوراق المالية', 'سوق المال', 'الاندماج', 'الاستحواذ', 'الأسهم', 'التداول']
        sama_match = any(kw in q_lower for kw in sama_en) or any(kw in query for kw in sama_ar)
        cma_match = any(kw in q_lower for kw in cma_en) or any(kw in query for kw in cma_ar)
        if sama_match and cma_match:
            return 'BOTH'
        elif sama_match:
            return 'SAMA'
        elif cma_match:
            return 'CMA'
        return 'BOTH'
    
    def translate_arabic_query(self, query):
        translations = {
            'رسوم الترخيص': 'licensing fees', 'رسوم ترخيص': 'licensing fees', 'شركات التمويل': 'finance companies',
            'شركة التمويل': 'finance company', 'التمويل العقاري': 'real estate finance', 'التمويل الأصغر': 'microfinance',
            'المستثمر المؤهل': 'qualified investor', 'مستثمر مؤهل': 'qualified investor', 'الصكوك': 'sukuk debt instruments',
            'صكوك': 'sukuk debt instruments', 'أدوات الدين': 'debt instruments', 'طرح عام': 'public offering',
            'طرح خاص': 'private placement', 'نشرة الإصدار': 'prospectus', 'صناديق الاستثمار': 'investment funds',
            'صندوق استثمار': 'investment fund', 'رأس المال': 'capital requirements', 'متطلبات رأس المال': 'capital requirements',
            'الحد الأدنى': 'minimum requirements', 'هيئة السوق المالية': 'capital market authority CMA',
            'مؤسسة النقد': 'SAMA monetary authority', 'ساما': 'SAMA', 'الاندماج': 'merger', 'الاستحواذ': 'acquisition',
            'الأسهم': 'shares stocks', 'الإفصاح': 'disclosure', 'الحوكمة': 'governance', 'مجلس الإدارة': 'board of directors',
            'تقرير سنوي': 'annual report', 'القوائم المالية': 'financial statements', 'المراجع الخارجي': 'external auditor',
            'العقوبات': 'penalties', 'المخالفات': 'violations', 'الترخيص': 'license licensing', 'التسجيل': 'registration',
            'الإدراج': 'listing', 'السوق الموازية': 'parallel market', 'الطرح': 'offering', 'الاكتتاب': 'subscription IPO',
        }
        result = query
        for ar, en in translations.items():
            if ar in query:
                result = result + ' ' + en
        return result if result != query else query
    
    def expand_query(self, query, lang):
        q = query.lower()
        mappings = {
            'sukuk': 'debt instruments securities bonds', 'sukuk issuance': 'debt instruments offering securities',
            'debt instruments': 'sukuk securities bonds', 'licensing fee': 'license fee financial consideration',
            'licensing fees': 'license fee financial consideration', 'qualified investor': 'accredited investor',
            'capital requirements': 'minimum capital paid up capital', 'finance company': 'finance companies',
            'microfinance': 'micro finance small finance', 'real estate finance': 'mortgage property finance',
            'investment fund': 'investment funds', 'public offering': 'IPO offering securities', 'private placement': 'exempt offering',
        }
        expansions = [exp for term, exp in mappings.items() if term in q]
        return query + ' ' + ' '.join(expansions) if expansions else query
    
    def bm25_search(self, query, regulator, language, top_k=15, force_english=False):
        search_lang = 'en' if force_english else language
        tokens = re.findall(r'[\u0600-\u06FF]+|[a-zA-Z]+|\d+', query.lower())
        if not tokens:
            return []
        scores = self.bm25.get_scores(tokens)
        results = []
        for idx in np.argsort(scores)[::-1]:
            if len(results) >= top_k or idx >= len(self.documents):
                continue
            doc = self.documents[idx]
            if regulator != 'BOTH' and doc.get('regulator') != regulator:
                continue
            if doc.get('language') != search_lang:
                continue
            if scores[idx] > 0:
                results.append({'doc': doc, 'score': float(scores[idx]), 'source': 'bm25'})
        return results
    
    def semantic_search(self, query, regulator, language, top_k=15, force_english=False):
        search_lang = 'en' if force_english else language
        embedding = self.embedder.encode([f"query: {query}"]).tolist()
        where = {"language": {"$eq": search_lang}} if regulator == 'BOTH' else {"$and": [{"language": {"$eq": search_lang}}, {"regulator": {"$eq": regulator}}]}
        try:
            results = self.collection.query(query_embeddings=embedding, n_results=top_k, where=where)
        except:
            results = self.collection.query(query_embeddings=embedding, n_results=top_k * 2)
        output = []
        if results['documents'] and results['documents'][0]:
            for doc_text, meta, dist in zip(results['documents'][0], results['metadatas'][0], results['distances'][0]):
                if regulator != 'BOTH' and meta.get('regulator') != regulator:
                    continue
                if meta.get('language') != search_lang:
                    continue
                output.append({'doc': {'text': doc_text, 'article': meta.get('article', ''), 'document': meta.get('document', ''), 'regulator': meta.get('regulator', ''), 'language': meta.get('language', '')}, 'score': 1/(1+dist), 'source': 'semantic'})
        return output[:top_k]
    
    def hybrid_search(self, query, n_results=3):
        user_language = self.detect_language(query)
        regulator = self.detect_regulator(query)
        if user_language == 'ar':
            english_query = self.translate_arabic_query(query)
            expanded = self.expand_query(english_query, 'en')
            print(f"DEBUG: Arabic Query → English Bridge")
            print(f"  Original: {query[:50]}")
            print(f"  Translated: {english_query[:50]}")
            print(f"  Regulator: {regulator}")
            bm25_res = self.bm25_search(expanded, regulator, user_language, force_english=True)
            sem_res = self.semantic_search(expanded, regulator, user_language, force_english=True)
        else:
            expanded = self.expand_query(query, user_language)
            print(f"DEBUG: English Query='{query[:50]}', Reg={regulator}")
            bm25_res = self.bm25_search(expanded, regulator, user_language)
            sem_res = self.semantic_search(expanded, regulator, user_language)
        print(f"DEBUG: BM25={len(bm25_res)}, Semantic={len(sem_res)}")
        doc_scores = {}
        k = 60
        for rank, r in enumerate(bm25_res):
            key = f"{r['doc']['document']}:{r['doc']['article']}"
            if key not in doc_scores:
                doc_scores[key] = {'doc': r['doc'], 'rrf': 0, 'src': set()}
            doc_scores[key]['rrf'] += 1/(k+rank+1)
            doc_scores[key]['src'].add('BM25')
        for rank, r in enumerate(sem_res):
            key = f"{r['doc']['document']}:{r['doc']['article']}"
            if key not in doc_scores:
                doc_scores[key] = {'doc': r['doc'], 'rrf': 0, 'src': set()}
            doc_scores[key]['rrf'] += 1/(k+rank+1)
            doc_scores[key]['src'].add('Semantic')
        sorted_res = sorted(doc_scores.values(), key=lambda x: x['rrf'], reverse=True)
        final = []
        for r in sorted_res[:n_results]:
            final.append(r['doc'])
            print(f"  → {r['doc']['article']} [{'+'.join(r['src'])}]")
        return final, regulator, user_language
    
    def is_follow_up(self, query):
        query_lower = query.lower().strip()
        follow_up_en = ['yes', 'yeah', 'sure', 'please', 'ok', 'okay', 'simplify', 'explain', 'example', 'examples', 'scenario', 'more details', 'elaborate', 'clarify', 'what do you mean', 'can you explain', 'help me understand', 'break it down', 'in simple terms', 'simpler', 'easier']
        follow_up_ar = ['نعم', 'أجل', 'طيب', 'حسنا', 'موافق', 'تمام', 'وضح', 'اشرح', 'مثال', 'أمثلة', 'سيناريو', 'تفاصيل أكثر', 'بسط', 'بشكل أبسط', 'ساعدني أفهم', 'ماذا تعني', 'اشرح أكثر']
        if any(pattern in query_lower for pattern in follow_up_en):
            return True
        if any(pattern in query for pattern in follow_up_ar):
            return True
        if len(query.strip()) < 15 and len(query.split()) <= 3:
            return True
        return False
    
    def is_out_of_domain(self, query):
        query_lower = query.lower()
        out_of_domain = ['weather', 'recipe', 'cook', 'movie', 'song', 'music', 'game', 'sport', 'football', 'soccer', 'basketball', 'joke', 'story', 'poem', 'write me', 'create a', 'translate', 'what is the capital', 'who is the president', 'how to code', 'python', 'javascript', 'programming', 'health', 'medical', 'doctor', 'disease', 'travel', 'hotel', 'flight', 'vacation', 'الطقس', 'وصفة', 'طبخ', 'فيلم', 'أغنية', 'موسيقى', 'لعبة', 'رياضة', 'كرة القدم', 'نكتة', 'قصة', 'قصيدة', 'ترجم', 'عاصمة', 'رئيس', 'برمجة', 'صحة', 'طبيب', 'سفر']
        return any(term in query_lower for term in out_of_domain)
    
    def build_prompt(self, question, docs, language, is_follow_up=False, conversation_context=None):
        ctx = "\n\n---\n\n".join([f"[Document {i}]\nSource: {d['document']}\nArticle: {d['article']}\nContent:\n{d['text']}" for i, d in enumerate(docs, 1)])
        conv_context = ""
        if is_follow_up and conversation_context:
            conv_context = "\n\nPrevious conversation:\n"
            for msg in conversation_context[-4:]:
                role = "User" if msg['role'] == 'user' else "Assistant"
                conv_context += f"{role}: {msg['content'][:500]}\n"
        if language == 'ar':
            if is_follow_up:
                return f"""أنت مساعد قانوني متخصص في الأنظمة المالية السعودية (ساما وهيئة السوق المالية).

المحادثة السابقة:
{conv_context}

المستندات المرجعية:
{ctx}

طلب المستخدم: {question}

المستخدم يطلب توضيحاً أو تبسيطاً. قم بما يلي:
- إذا طلب تبسيط: اشرح المفهوم بلغة سهلة وواضحة
- إذا طلب مثال: قدم سيناريو عملي يوضح التطبيق
- إذا طلب توضيح: اشرح النقاط الغامضة بالتفصيل

الإجابة:"""
            else:
                return f"""أنت مساعد قانوني متخصص في الأنظمة المالية السعودية (ساما وهيئة السوق المالية).

تعليمات مهمة:
- اقرأ كل مستند بعناية قبل الإجابة
- استخرج المعلومات المطلوبة من المستندات
- اذكر رقم المادة عند الاستشهاد (مثال: المادة 22)
- مصطلح "debt instruments" يعني "الصكوك" أو "أدوات الدين"
- اكتب الأرقام والمبالغ كما هي في المستند
- استخدم تنسيق Markdown: استخدم **نص** للتأكيد و - للقوائم
- إذا لم تجد المعلومة المحددة، قل ذلك بوضوح
- لا تذكر أنك ستساعد في نهاية الإجابة

المستندات المرجعية:
{ctx}

السؤال: {question}

الإجابة:"""
        if is_follow_up:
            return f"""You are a legal assistant specializing in Saudi Arabian financial regulations (SAMA and CMA).

Previous conversation:
{conv_context}

Reference Documents:
{ctx}

User request: {question}

The user is asking for clarification or simplification. Do the following:
- If they want simplification: Explain the concept in plain, easy-to-understand language
- If they want examples: Provide a practical scenario showing how this applies
- If they want clarification: Explain the unclear points in detail

Answer:"""
        else:
            return f"""You are a legal assistant specializing in Saudi Arabian financial regulations (SAMA and CMA).

Important instructions:
- Read each document carefully before answering
- Extract the relevant information from the documents
- Cite the Article number when referencing information (e.g., "Article 22")
- "Sukuk" and "debt instruments" refer to the same thing
- Preserve exact numbers and amounts as written in the documents
- Use Markdown formatting: **bold** for emphasis and - for lists
- If the specific information is not found, say so clearly
- Do not offer further assistance at the end of your response

Reference Documents:
{ctx}

Question: {question}

Answer:"""
    
    def build_out_of_domain_response(self, language):
        if language == 'ar':
            return """عذراً، أنا مساعد متخصص في الأنظمة المالية السعودية فقط.

يمكنني مساعدتك في:
- **أنظمة ساما**: شركات التمويل، التمويل العقاري، التمويل الأصغر
- **أنظمة هيئة السوق المالية**: الصكوك، صناديق الاستثمار، المستثمر المؤهل، الطرح والإدراج

يرجى طرح سؤال يتعلق بهذه المواضيع."""
        else:
            return """I apologize, but I am a specialized assistant for Saudi Arabian financial regulations only.

I can help you with:
- **SAMA regulations**: Finance companies, real estate finance, microfinance
- **CMA regulations**: Sukuk, investment funds, qualified investors, offerings and listings

Please ask a question related to these topics."""
    
    def _chunk_text(self, text, chunk_size=1000, overlap=200):
        """Helper to split uploaded text into overlap chunks for analysis."""
        if not text: return []
        chunks = []
        start = 0
        text_len = len(text)
        while start < text_len:
            end = start + chunk_size
            chunks.append(text[start:end])
            start += (chunk_size - overlap)
        return chunks

    def generate_executive_brief(self):
        """
        Generates a 3-part Executive Summary using manual chunking and on-the-fly embedding.
        Corrected to match TadqeeqRAG v3.0 architecture.
        """
        import numpy as np
        import ollama

        # 1. Get text from the uploaded document processor
        text = self.doc_processor.get_current_text()
        if not text:
            return {"error": "No document uploaded. Please upload a file first."}

        print("--- Starting Executive Brief Generation ---")

        # 2. Manual Chunking (Split text into manageable pieces)
        chunk_size = 1000
        overlap = 200
        chunks = []
        start = 0
        text_len = len(text)
        
        if text_len < 50:
            return {"error": "Document is too short to analyze."}

        while start < text_len:
            end = start + chunk_size
            chunks.append(text[start:end])
            start += (chunk_size - overlap)
        
        print(f"--- Document split into {len(chunks)} chunks ---")

        # 3. Embed chunks (On-the-fly)
        try:
            chunk_embeddings = self.embedder.encode(chunks)
        except Exception as e:
            return {"error": f"Embedding failed: {str(e)}"}

        # 4. Multi-Targeted Retrieval
        targets = [
            "What are the key risks, violations, penalties, and compliance red flags?",
            "What are the fees, capital requirements, costs, and financial obligations?",
            "What are the effective dates, deadlines, submission timelines, and expiry dates?"
        ]

        unique_indices = set()
        
        for query in targets:
            query_embedding = self.embedder.encode(query)
            # Calculate Cosine Similarity via Dot Product
            scores = np.dot(chunk_embeddings, query_embedding)
            # Get top 5 indices for this query
            top_k_indices = np.argsort(scores)[-5:][::-1]
            for idx in top_k_indices:
                unique_indices.add(idx)

        # 5. Assemble Context
        relevant_chunks = [chunks[i] for i in unique_indices]
        combined_context = "\n\n...\n\n".join(relevant_chunks)
        print(f"--- Analyzed {len(relevant_chunks)} unique chunks ---")

        # 6. Prompt
        prompt = f"""You are a Senior Risk Analyst. Your job is to write a high-level "Executive Brief" based ONLY on the text provided below.
        
        Strictly follow this Markdown structure:

        # 📋 Executive Brief

        ## 🚨 Key Risks & Red Flags
        * (List the most critical compliance risks, prohibitions, or penalties found. Be concise.)

        ## 💰 Financial Obligations
        * (List specific fees, capital requirements, fines, or recurring costs. Use numbers if available.)

        ## 📅 Critical Deadlines
        * (List specific dates, submission timelines, or effective periods mentioned.)

        If a section has no relevant information in the text, write "No significant details found."

        CONTEXT DATA:
        {combined_context}
        """

        # 7. Generate via Ollama
        try:
            # We use the globally defined LLM_MODEL from your config
            resp = ollama.generate(model=LLM_MODEL, prompt=prompt, options={'temperature': 0.1})
            result = resp['response'].strip()
            self.last_brief = result
            return {"report": result}
        except Exception as e:
            print(f"Error generating brief: {e}")
            return {"error": f"Failed to generate brief: {str(e)}"}
    
    def generate_response(self, question):
        import ollama
        lang = self.detect_language(question)
        if self.is_out_of_domain(question):
            return {'answer': self.build_out_of_domain_response(lang), 'sources': [], 'regulator': 'NONE'}
        is_followup = self.is_follow_up(question)
        conversation_context = None
        if is_followup:
            conversation_context = self.chat_history.get_conversation_context()
        docs, reg, lang = self.hybrid_search(question)
        if not docs:
            no_info = 'No relevant information found.' if lang == 'en' else 'لم يتم العثور على معلومات ذات صلة.'
            return {'answer': no_info, 'sources': [], 'regulator': reg}
        prompt = self.build_prompt(question, docs, lang, is_followup, conversation_context)
        resp = ollama.generate(model=LLM_MODEL, prompt=prompt, options={'temperature': 0.1, 'num_predict': 2000})
        seen = set()
        sources = [{'article': d['article'], 'document': d['document']} for d in docs if d['article'] not in seen and not seen.add(d['article'])]
        return {'answer': resp['response'].strip(), 'sources': sources, 'regulator': reg}


# ==============================================================================
# API CLASS
# ==============================================================================

class API:

    def __init__(self):
        self.rag = None
        self.window = None
        self._init_thread = None
        self._init_started = False
        self.last_compliance_result = None

    def generate_brief(self):
        """Endpoint for the UI to trigger the Executive Brief generation."""
        if not self.rag:
            return {'error': 'System not initialized'}
        return self.rag.generate_executive_brief()
    
    
    def set_window(self, window):
        self.window = window
    
    def get_init_status(self):
        global progress_tracker
        status = progress_tracker.get_status()
        if TadqeeqRAG._initialized and self.rag:
            return {
                'status': 'ready', 'progress': 100, 'stage': 'ready', 'stage_text': 'Ready!',
                'total': self.rag.total, 'sama': self.rag.sama_count, 'cma': self.rag.cma_count,
                'sama_en': self.rag.stats['SAMA']['en'], 'sama_ar': self.rag.stats['SAMA']['ar'],
                'cma_en': self.rag.stats['CMA']['en'], 'cma_ar': self.rag.stats['CMA']['ar'],
                'chats': self.rag.chat_history.get_recent_chats()
            }
        if TadqeeqRAG._init_error:
            return {'status': 'error', 'error': TadqeeqRAG._init_error, **status}
        return {'status': 'loading', **status}
    
    def start_initialization(self):
        if self._init_started:
            return {'status': 'already_started'}
        self._init_started = True
        def init_background():
            self.rag = TadqeeqRAG.get_instance()
            self.rag.initialize()
        self._init_thread = threading.Thread(target=init_background, daemon=True)
        self._init_thread.start()
        return {'status': 'started'}
    
    def initialize(self):
        if not self._init_started:
            self.start_initialization()
        timeout = 120
        start_time = time.time()
        while time.time() - start_time < timeout:
            if TadqeeqRAG._initialized:
                return self.get_init_status()
            if TadqeeqRAG._init_error:
                return self.get_init_status()
            time.sleep(0.5)
        return {'status': 'error', 'message': 'Initialization timeout'}
    
    def upload_document(self, file_data, filename):
        if not self.rag:
            return {'error': 'System not initialized'}
        return self.rag.doc_processor.process_file(file_data=file_data, filename=filename)
    
    def run_compliance_check(self):
        if not self.rag:
            return {'error': 'System not initialized'}
        text = self.rag.doc_processor.get_current_text()
        if not text:
            return {'error': 'No document uploaded'}
        
        # 1. Run the check
        final_result = self.rag.compliance_checker.check_compliance(text, self.rag.doc_processor.current_filename)
        
        # 2. SAVE THE RESULT (This is what you were missing!)
        self.last_compliance_result = final_result 
        
        return final_result
    
    def clear_document(self):
        if not self.rag:
            return {'error': 'System not initialized'}
        self.rag.doc_processor.clear()
        return {'success': True}
    
    def export_markdown(self):
        try:
            import webview
            if not self.rag:
                return {'error': 'System not initialized'}
            messages = self.rag.chat_history.get_messages()
            if not messages:
                return {'error': 'No messages to export. Start a conversation first.'}
            md_content, error = self.rag.chat_exporter.export_markdown(messages)
            if error:
                return {'error': error}
            filename = f'tadqeeq_chat_{datetime.now().strftime("%Y%m%d_%H%M%S")}.md'
            if self.window:
                try:
                    save_path = self.window.create_file_dialog(webview.SAVE_DIALOG, save_filename=filename, file_types=('Markdown Files (*.md)',))
                except:
                    save_path = self.window.create_file_dialog(dialog_type=webview.SAVE_DIALOG, save_filename=filename, file_types=('Markdown Files (*.md)',))
                if save_path:
                    path = save_path if isinstance(save_path, str) else save_path[0] if save_path else None
                    if path:
                        with open(path, 'w', encoding='utf-8') as f:
                            f.write(md_content)
                        return {'success': True, 'path': path}
                return {'error': 'Export cancelled'}
            else:
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(md_content)
                return {'success': True, 'path': filename}
        except Exception as e:
            return {'error': str(e)}
    
    def export_pdf(self):
        try:
            import webview
            if not self.rag:
                return {'error': 'System not initialized'}
            messages = self.rag.chat_history.get_messages()
            if not messages:
                return {'error': 'No messages to export. Start a conversation first.'}
            pdf_data, error = self.rag.chat_exporter.export_pdf(messages)
            if error:
                return {'error': error}
            pdf_bytes = base64.b64decode(pdf_data)
            filename = f'tadqeeq_chat_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
            if self.window:
                try:
                    save_path = self.window.create_file_dialog(webview.SAVE_DIALOG, save_filename=filename, file_types=('PDF Files (*.pdf)',))
                except:
                    save_path = self.window.create_file_dialog(dialog_type=webview.SAVE_DIALOG, save_filename=filename, file_types=('PDF Files (*.pdf)',))
                if save_path:
                    path = save_path if isinstance(save_path, str) else save_path[0] if save_path else None
                    if path:
                        with open(path, 'wb') as f:
                            f.write(pdf_bytes)
                        return {'success': True, 'path': path}
                return {'error': 'Export cancelled'}
            else:
                with open(filename, 'wb') as f:
                    f.write(pdf_bytes)
                return {'success': True, 'path': filename}
        except Exception as e:
            return {'error': str(e)}
    
    def query(self, question):
        if not self.rag:
            return {'error': 'System not initialized'}
        self.rag.chat_history.add_message('user', question)
        result = self.rag.generate_response(question)
        self.rag.chat_history.add_message('assistant', result['answer'], result.get('sources'), result.get('regulator'))
        return result
    
    def new_chat(self):
        if not self.rag:
            return {'error': 'System not initialized'}
        self.rag.doc_processor.clear()
        chat_id = self.rag.chat_history.new_chat()
        return {'id': chat_id, 'chats': self.rag.chat_history.get_recent_chats()}
    
    def load_chat(self, chat_id):
        if not self.rag:
            return {'error': 'System not initialized'}
        messages = self.rag.chat_history.load_chat(chat_id)
        return {'messages': messages}
    
    def get_chats(self):
        if not self.rag:
            return {'error': 'System not initialized'}
        return {'chats': self.rag.chat_history.get_recent_chats()}
    
    def delete_chat(self, chat_id):
        if not self.rag:
            return {'error': 'System not initialized'}
        success = self.rag.chat_history.delete_chat(chat_id)
        return {'success': success, 'chats': self.rag.chat_history.get_recent_chats()}
    
    def export_brief_markdown(self):
        import webview
        if not self.rag or not self.rag.last_brief: return {'error': 'No brief generated'}
        
        filename = f'Executive_Brief_{datetime.now().strftime("%Y%m%d_%H%M%S")}.md'
        content, err = self.rag.chat_exporter.export_brief_markdown(self.rag.last_brief)
        
        if self.window:
             try:
                 save_path = self.window.create_file_dialog(webview.SAVE_DIALOG, save_filename=filename, file_types=('Markdown Files (*.md)',))
             except:
                 save_path = self.window.create_file_dialog(dialog_type=webview.SAVE_DIALOG, save_filename=filename, file_types=('Markdown Files (*.md)',))
             
             if save_path:
                 path = save_path if isinstance(save_path, str) else save_path[0] if save_path else None
                 if path:
                     with open(path, 'w', encoding='utf-8') as f: f.write(content)
                     return {'success': True, 'path': path}
        return {'error': 'Export cancelled'}

    def export_brief_pdf(self):
        import webview
        if not self.rag or not self.rag.last_brief: return {'error': 'No brief generated'}
        
        filename = f'Executive_Brief_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
        content_b64, err = self.rag.chat_exporter.export_brief_pdf(self.rag.last_brief)
        if err: return {'error': err}
        
        if self.window:
             try:
                 save_path = self.window.create_file_dialog(webview.SAVE_DIALOG, save_filename=filename, file_types=('PDF Files (*.pdf)',))
             except:
                 save_path = self.window.create_file_dialog(dialog_type=webview.SAVE_DIALOG, save_filename=filename, file_types=('PDF Files (*.pdf)',))
             
             if save_path:
                 path = save_path if isinstance(save_path, str) else save_path[0] if save_path else None
                 if path:
                     with open(path, 'wb') as f: f.write(base64.b64decode(content_b64))
                     return {'success': True, 'path': path}
        return {'error': 'Export cancelled'}
    
    def export_compliance_pdf(self):
        """Exports the compliance report with a clean, single-table layout."""
        import webview
        if not self.last_compliance_result:
            return {"error": "No compliance report available to export."}
        
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        
        try:
            data = self.last_compliance_result
            filename = f"Compliance_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            
            buffer = BytesIO()
            # 17cm usable width for A4
            doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
            styles = getSampleStyleSheet()
            story = []

            # 1. Professional Header
            title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=22, alignment=0, spaceAfter=10, textColor=colors.HexColor('#00d4aa'))
            story.append(Paragraph("Regulatory Compliance Audit", title_style))
            story.append(Paragraph(f"<b>Document:</b> {data['filename']}", styles['Normal']))
            story.append(Paragraph(f"<b>Date:</b> {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
            story.append(Paragraph(f"<b>Final Score:</b> {data['score']}% Compliant", styles['Normal']))
            story.append(Spacer(1, 1.5*cm))

            # 2. Detailed Checklist Table
            table_data = [['Status', 'Regulation / Check', 'Findings & Details']]
            for check in data['checks']:
                status_text = "PASS" if check['status'] == 'compliant' else "WARN"
                table_data.append([
                    status_text, 
                    Paragraph(f"<b>{check['regulation']}</b>", styles['Normal']), 
                    Paragraph(check['detail'], styles['Normal'])
                ])

            checklist_table = Table(table_data, colWidths=[2*cm, 6*cm, 9*cm])
            checklist_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#0f172a')), # Navy Header
                ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('GRID', (0,0), (-1,-1), 0.5, colors.lightgrey),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('TOPPADDING', (0,0), (-1,-1), 12),
                ('BOTTOMPADDING', (0,0), (-1,-1), 12),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,0), 11),
            ]))
            story.append(checklist_table)

            # Build and Export
            doc.build(story)
            pdf_b64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
            buffer.close()

            if self.window:
                save_path = self.window.create_file_dialog(webview.SAVE_DIALOG, save_filename=filename, file_types=('PDF Files (*.pdf)',))
                if save_path:
                    path = save_path if isinstance(save_path, str) else save_path[0]
                    with open(path, 'wb') as f:
                        f.write(base64.b64decode(pdf_b64))
                    return {'success': True, 'path': path}
            return {'error': 'Export cancelled'}

        except Exception as e:
            return {"error": f"Compliance PDF Export Error: {str(e)}"}