"""
TadqeeqAI v2.2
==============
Bilingual RAG system for Saudi Arabian financial regulations.
SAMA + CMA · English + Arabic

Features:
- Hybrid search: BM25 + Semantic (multilingual-e5-base)
- Chat history with persistent storage
- Follow-up support (simplify, examples)
- Domain-restricted responses
- Auto-start Ollama (hidden, auto-shutdown)

v2.2 Additions:
- Document upload & analysis (PDF, DOCX)
- Regulatory compliance checker
- Chat export (Markdown, PDF)
- Drag & drop file upload
- Gemini-style input bar

Prerequisites:
    1. Ensure chroma_db_v2 and bm25_index.pkl exist
    2. Install: pip install PyMuPDF python-docx reportlab
    3. Ollama with aya:8b model
"""

import json
import pickle
import re
import os
import subprocess
import time
import uuid
import warnings
import logging
from datetime import datetime
from pathlib import Path

# Suppress pywebview accessibility warnings
warnings.filterwarnings('ignore', category=DeprecationWarning)
logging.getLogger('pywebview').setLevel(logging.ERROR)

import chromadb
from sentence_transformers import SentenceTransformer
from rank_bm25 import BM25Okapi
import ollama
import numpy as np
import webview

import base64
import tempfile
from io import BytesIO

# Document processing (install: pip install PyMuPDF python-docx reportlab)
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_CENTER
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

# Configuration
CHROMA_PATH = "./chroma_db_v2"
BM25_PATH = "./bm25_index.pkl"
DOCS_PATH = "./documents.json"
CHAT_HISTORY_PATH = "./chat_history"
EMBEDDING_MODEL = 'intfloat/multilingual-e5-base'
MAX_DOCUMENT_PAGES = 50

# LLM Model Options:
# - 'aya:8b'       : Best for Arabic, no Chinese leak, follows instructions well
# - 'qwen2.5:7b'   : Good overall but leaks Chinese on long Arabic responses  
# - 'jais:13b'     : Arabic-first but truncates responses
LLM_MODEL = 'aya:8b'

# Ensure chat history directory exists
Path(CHAT_HISTORY_PATH).mkdir(exist_ok=True)

# Global Ollama process reference
OLLAMA_PROCESS = None


def start_ollama():
    """Start Ollama server hidden, preventing GUI app from launching."""
    import atexit
    global OLLAMA_PROCESS
    
    print("  Checking Ollama...")
    
    if os.name == 'nt':
        # Windows: Kill ALL Ollama processes first (GUI app, server, everything)
        # This ensures we start fresh with our hidden instance
        for proc_name in ['ollama app.exe', 'Ollama.exe', 'ollama.exe']:
            subprocess.run(
                ['taskkill', '/F', '/IM', proc_name],
                capture_output=True,
                creationflags=subprocess.CREATE_NO_WINDOW
            )
        time.sleep(2)  # Wait for processes to fully terminate
    
    print("  Starting Ollama (hidden)...")
    try:
        if os.name == 'nt':
            # Windows: Start ollama serve with completely hidden window
            # Use shell=False and specify full environment to prevent GUI trigger
            startupinfo = subprocess.STARTUPINFO()
            startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW | subprocess.STARTF_USESTDHANDLES
            startupinfo.wShowWindow = 0  # SW_HIDE
            
            # Find ollama.exe path (usually in user's AppData or Program Files)
            ollama_paths = [
                os.path.expandvars(r'%LOCALAPPDATA%\Programs\Ollama\ollama.exe'),
                os.path.expandvars(r'%LOCALAPPDATA%\Ollama\ollama.exe'),
                r'C:\Program Files\Ollama\ollama.exe',
                'ollama'  # Fallback to PATH
            ]
            
            ollama_exe = 'ollama'
            for path in ollama_paths:
                if os.path.exists(path):
                    ollama_exe = path
                    break
            
            OLLAMA_PROCESS = subprocess.Popen(
                [ollama_exe, 'serve'],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                stdin=subprocess.DEVNULL,
                startupinfo=startupinfo,
                creationflags=subprocess.CREATE_NO_WINDOW | subprocess.DETACHED_PROCESS | subprocess.CREATE_NEW_PROCESS_GROUP
            )
        else:
            # Linux/Mac
            OLLAMA_PROCESS = subprocess.Popen(
                ['ollama', 'serve'],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL
            )
        
        # Register cleanup
        atexit.register(stop_ollama)
        
        # Wait for server to be ready
        for i in range(30):
            time.sleep(1)
            try:
                # Use requests or simple socket check instead of subprocess
                # to avoid triggering GUI
                import socket
                sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                result = sock.connect_ex(('127.0.0.1', 11434))
                sock.close()
                if result == 0:
                    print("  ✓ Ollama started (hidden)")
                    return True
            except:
                pass
            if i % 10 == 9:
                print(f"    Waiting... ({i+1}s)")
        
        print("  ⚠ Ollama may not have started - continuing anyway")
        return False
    except Exception as e:
        print(f"  ⚠ Could not start Ollama: {e}")
        return False


def stop_ollama():
    """Stop Ollama when app closes."""
    global OLLAMA_PROCESS
    if OLLAMA_PROCESS:
        try:
            OLLAMA_PROCESS.terminate()
            OLLAMA_PROCESS.wait(timeout=5)
        except:
            try:
                OLLAMA_PROCESS.kill()
            except:
                pass
        OLLAMA_PROCESS = None
    
    # Also kill any remaining ollama processes on Windows
    if os.name == 'nt':
        for proc_name in ['ollama.exe']:
            subprocess.run(
                ['taskkill', '/F', '/IM', proc_name],
                capture_output=True,
                creationflags=subprocess.CREATE_NO_WINDOW
            )


class ChatHistory:
    """Manages persistent chat history."""
    
    def get_messages(self):
        return self.current_messages

    def __init__(self):
        self.history_path = Path(CHAT_HISTORY_PATH)
        self.current_chat_id = None
        self.current_messages = []
    
    def new_chat(self):
        """Start a new chat session."""
        self.current_chat_id = str(uuid.uuid4())[:8]
        self.current_messages = []
        return self.current_chat_id
    
    def add_message(self, role, content, sources=None, regulator=None):
        """Add a message to current chat."""
        msg = {
            'role': role,
            'content': content,
            'timestamp': datetime.now().isoformat()
        }
        if sources:
            msg['sources'] = sources
        if regulator:
            msg['regulator'] = regulator
        self.current_messages.append(msg)
        self._save_current()
    
    def _save_current(self):
        """Save current chat to file."""
        if not self.current_chat_id:
            return
        filepath = self.history_path / f"{self.current_chat_id}.json"
        chat_data = {
            'id': self.current_chat_id,
            'created': self.current_messages[0]['timestamp'] if self.current_messages else datetime.now().isoformat(),
            'updated': datetime.now().isoformat(),
            'messages': self.current_messages,
            'preview': self._get_preview()
        }
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(chat_data, f, ensure_ascii=False, indent=2)
    
    def _get_preview(self):
        """Get first user message as preview."""
        for msg in self.current_messages:
            if msg['role'] == 'user':
                text = msg['content']
                return text[:50] + '...' if len(text) > 50 else text
        return 'New Chat'
    
    def get_recent_chats(self, limit=20):
        """Get list of recent chats."""
        chats = []
        for filepath in sorted(self.history_path.glob('*.json'), key=lambda x: x.stat().st_mtime, reverse=True):
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    chats.append({
                        'id': data['id'],
                        'preview': data.get('preview', 'Chat'),
                        'updated': data.get('updated', ''),
                        'message_count': len(data.get('messages', []))
                    })
            except:
                pass
            if len(chats) >= limit:
                break
        return chats
    
    def load_chat(self, chat_id):
        """Load a specific chat."""
        filepath = self.history_path / f"{chat_id}.json"
        if filepath.exists():
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
                self.current_chat_id = chat_id
                self.current_messages = data.get('messages', [])
                return self.current_messages
        return []
    
    def get_conversation_context(self, limit=6):
        """Get recent messages for context (for follow-ups)."""
        return self.current_messages[-limit:] if self.current_messages else []
    
    def delete_chat(self, chat_id):
        """Delete a chat permanently."""
        filepath = self.history_path / f"{chat_id}.json"
        if filepath.exists():
            filepath.unlink()
            # If we deleted the current chat, reset
            if self.current_chat_id == chat_id:
                self.current_chat_id = None
                self.current_messages = []
            return True
        return False
    
class DocumentProcessor:
    """Process uploaded documents (PDF, DOCX)."""
    
    def __init__(self):
        self.current_document = None
        self.current_text = None
        self.current_filename = None
        self.current_page_count = 0
    
    def process_file(self, file_path=None, file_data=None, filename=None):
        """Process a document file. Accepts either file path or base64 data."""
        try:
            if file_data:
                data = base64.b64decode(file_data)
                ext = filename.lower().split('.')[-1] if filename else ''
            elif file_path:
                with open(file_path, 'rb') as f:
                    data = f.read()
                ext = file_path.lower().split('.')[-1]
                filename = os.path.basename(file_path)
            else:
                return {'error': 'No file provided'}
            
            if ext == 'pdf':
                return self._process_pdf(data, filename)
            elif ext in ['docx', 'doc']:
                return self._process_docx(data, filename)
            else:
                return {'error': f'Unsupported file type: {ext}'}
                
        except Exception as e:
            return {'error': str(e)}
    
    def _process_pdf(self, data, filename):
        if not HAS_PYMUPDF:
            return {'error': 'PDF support not available. Install PyMuPDF: pip install PyMuPDF'}
        
        try:
            doc = fitz.open(stream=data, filetype="pdf")
            
            if doc.page_count > MAX_DOCUMENT_PAGES:
                return {'error': f'Document too large. Maximum {MAX_DOCUMENT_PAGES} pages allowed.'}
            
            text_parts = []
            for page_num, page in enumerate(doc):
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")
            
            full_text = "\n\n".join(text_parts)
            
            self.current_document = doc
            self.current_text = full_text
            self.current_filename = filename
            self.current_page_count = doc.page_count
            
            summary = self._generate_quick_summary(full_text, filename)
            
            return {
                'success': True,
                'filename': filename,
                'pages': doc.page_count,
                'chars': len(full_text),
                'summary': summary
            }
            
        except Exception as e:
            return {'error': f'Failed to process PDF: {str(e)}'}
    
    def _process_docx(self, data, filename):
        if not HAS_DOCX:
            return {'error': 'DOCX support not available. Install python-docx: pip install python-docx'}
        
        try:
            doc = DocxDocument(BytesIO(data))
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            full_text = "\n\n".join(text_parts)
            
            self.current_text = full_text
            self.current_filename = filename
            self.current_page_count = len(text_parts) // 20 + 1
            
            summary = self._generate_quick_summary(full_text, filename)
            
            return {
                'success': True,
                'filename': filename,
                'pages': self.current_page_count,
                'chars': len(full_text),
                'summary': summary
            }
            
        except Exception as e:
            return {'error': f'Failed to process DOCX: {str(e)}'}
    
    def _generate_quick_summary(self, text, filename):
        text_lower = text.lower()
        
        doc_type = "Document"
        if any(kw in text_lower for kw in ['fund', 'investment', 'investor', 'subscription']):
            doc_type = "Investment Fund Document"
        elif any(kw in text_lower for kw in ['sukuk', 'bond', 'debt instrument']):
            doc_type = "Sukuk/Debt Instrument Document"
        elif any(kw in text_lower for kw in ['license', 'licensing', 'authorization']):
            doc_type = "Licensing Document"
        elif any(kw in text_lower for kw in ['contract', 'agreement', 'party', 'parties']):
            doc_type = "Contract/Agreement"
        elif any(kw in text_lower for kw in ['prospectus', 'offering', 'securities']):
            doc_type = "Securities Prospectus"
        
        has_arabic = bool(re.search(r'[\u0600-\u06FF]', text))
        
        return {'type': doc_type, 'has_arabic': has_arabic, 'word_count': len(text.split())}
    
    def get_current_text(self):
        return self.current_text
    
    def clear(self):
        self.current_document = None
        self.current_text = None
        self.current_filename = None
        self.current_page_count = 0


class ComplianceChecker:
    """Check documents for regulatory compliance."""
    
    COMPLIANCE_CATEGORIES = {
        'qualified_investor': {
            'name': 'Qualified Investor Definition',
            'keywords': ['qualified investor', 'accredited investor', 'مستثمر مؤهل'],
            'regulation': 'CMA Rules on Offer of Securities, Article 15',
            'description': 'Documents offering securities must define qualified investor criteria'
        },
        'risk_disclosure': {
            'name': 'Risk Disclosure',
            'keywords': ['risk', 'risks', 'risk factors', 'مخاطر'],
            'regulation': 'CMA Rules on Offer of Securities, Article 22',
            'description': 'Offering documents must include comprehensive risk disclosures'
        },
        'capital_requirements': {
            'name': 'Capital Requirements',
            'keywords': ['minimum capital', 'paid-up capital', 'رأس المال'],
            'regulation': 'Finance Companies Control Law, Article 5',
            'description': 'Finance companies must meet minimum capital requirements'
        },
        'license_reference': {
            'name': 'Licensing Information',
            'keywords': ['license', 'licensed', 'CMA', 'SAMA', 'ترخيص'],
            'regulation': 'Capital Market Law, Article 3',
            'description': 'Financial activities require proper licensing'
        },
        'fund_terms': {
            'name': 'Fund Terms & Conditions',
            'keywords': ['terms and conditions', 'subscription', 'management fee', 'شروط وأحكام'],
            'regulation': 'Investment Funds Regulations, Article 20',
            'description': 'Investment funds must clearly state terms and fees'
        },
        'disclosure_requirements': {
            'name': 'Disclosure Requirements',
            'keywords': ['disclosure', 'material information', 'إفصاح'],
            'regulation': 'CMA Rules on Offer of Securities, Article 30',
            'description': 'Issuers must disclose all material information'
        }
    }
    
    def check_compliance(self, text, filename=None):
        results = {
            'filename': filename or 'Document',
            'timestamp': datetime.now().isoformat(),
            'checks': [],
            'summary': {'compliant': 0, 'warnings': 0, 'missing': 0}
        }
        
        text_lower = text.lower()
        
        for category_id, category in self.COMPLIANCE_CATEGORIES.items():
            found_keywords = [kw for kw in category['keywords'] if kw.lower() in text_lower or kw in text]
            
            if found_keywords:
                status = 'compliant'
                results['summary']['compliant'] += 1
                detail = f"Found references: {', '.join(found_keywords[:3])}"
            else:
                status = 'warning'
                results['summary']['warnings'] += 1
                detail = f"Consider adding {category['name'].lower()} information"
            
            results['checks'].append({
                'id': category_id,
                'name': category['name'],
                'status': status,
                'regulation': category['regulation'],
                'description': category['description'],
                'detail': detail
            })
        
        total = results['summary']['compliant'] + results['summary']['warnings']
        results['score'] = round((results['summary']['compliant'] / total) * 100) if total > 0 else 100
        
        return results


class ChatExporter:
    """Export chat conversations to Markdown and PDF."""
    
    def _sanitize_for_pdf(self, text):
        """Clean text for ReportLab PDF - remove markdown and escape XML."""
        import html
        # First escape any existing HTML/XML entities
        text = html.escape(text)
        # Remove markdown bold
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        # Remove markdown italic
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        # Remove markdown headers
        text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
        # Convert bullet points to dashes
        text = re.sub(r'^[\-\*]\s+', '- ', text, flags=re.MULTILINE)
        # Remove any remaining problematic characters
        text = text.replace('<', '(').replace('>', ')')
        return text
    
    def export_markdown(self, messages, filename=None):
        if not messages:
            return None, "No messages to export"
        
        md_lines = [
            "# TadqeeqAI Chat Export",
            f"**Date:** {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            "", "---", ""
        ]
        
        for msg in messages:
            role = "User" if msg['role'] == 'user' else "TadqeeqAI"
            md_lines.extend([f"## {role}", "", msg['content'], ""])
            
            if msg.get('sources'):
                md_lines.append("**Sources:**")
                for src in msg['sources']:
                    md_lines.append(f"- {src['article']} ({src['document']})")
                md_lines.append("")
            
            md_lines.extend(["---", ""])
        
        md_lines.extend(["", "*Generated by TadqeeqAI v2.2*"])
        return "\n".join(md_lines), None
    
    def export_pdf(self, messages, filename=None):
        if not messages:
            return None, "No messages to export"
        
        if not HAS_REPORTLAB:
            return None, "PDF export not available. Install reportlab: pip install reportlab"
        
        try:
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
            
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=24, spaceAfter=20, textColor=colors.HexColor('#00d4aa'))
            normal_style = ParagraphStyle('CustomNormal', parent=styles['Normal'], fontSize=11, leading=16)
            role_style = ParagraphStyle('CustomRole', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#00d4aa'), fontName='Helvetica-Bold', spaceBefore=15)
            footer_style = ParagraphStyle('CustomFooter', parent=styles['Normal'], fontSize=8, textColor=colors.grey, alignment=TA_CENTER)
            
            story = [
                Paragraph("TadqeeqAI Chat Export", title_style),
                Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}", styles['Normal']),
                HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e0e0e0')),
                Spacer(1, 20)
            ]
            
            for msg in messages:
                role = "You" if msg['role'] == 'user' else "TadqeeqAI"
                story.append(Paragraph(role, role_style))
                
                # Sanitize content for PDF
                content = self._sanitize_for_pdf(msg['content'])
                # Split into paragraphs and add each
                paragraphs = content.split('\n')
                for para in paragraphs:
                    para = para.strip()
                    if para:
                        try:
                            story.append(Paragraph(para, normal_style))
                        except:
                            # If paragraph still fails, use plain text
                            story.append(Paragraph(para.encode('ascii', 'ignore').decode(), normal_style))
                
                story.append(Spacer(1, 10))
                
                # Add sources if present
                if msg.get('sources'):
                    sources_text = "Sources: " + ", ".join([s['article'] for s in msg['sources'][:3]])
                    story.append(Paragraph(sources_text, ParagraphStyle('Sources', parent=styles['Normal'], fontSize=9, textColor=colors.grey)))
                    story.append(Spacer(1, 5))
            
            story.extend([
                Spacer(1, 30), 
                HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e0e0e0')), 
                Spacer(1, 10), 
                Paragraph("Generated by TadqeeqAI v2.2", footer_style)
            ])
            
            doc.build(story)
            pdf_data = buffer.getvalue()
            buffer.close()
            
            return base64.b64encode(pdf_data).decode('utf-8'), None
            
        except Exception as e:
            return None, f"PDF export failed: {str(e)}"


class TadqeeqRAG:
    """Hybrid RAG system with BM25 + semantic search."""
    _instance = None
    
    @classmethod
    def get_instance(cls):
        if cls._instance is None:
            cls._instance = cls()
        return cls._instance
    
    def __init__(self):
        print("Loading TadqeeqAI v2.2...")
        
        # Start Ollama first
        print("  Step 1: Starting Ollama...")
        start_ollama()
        
        # Load documents
        print("  Step 2: Loading documents...")
        with open(DOCS_PATH, 'r', encoding='utf-8') as f:
            self.documents = json.load(f)
        
        # Count by regulator/language
        self.stats = {'SAMA': {'en': 0, 'ar': 0}, 'CMA': {'en': 0, 'ar': 0}}
        for doc in self.documents:
            reg = doc.get('regulator', 'CMA')
            lang = doc.get('language', 'en')
            self.stats[reg][lang] += 1
        
        self.sama_count = self.stats['SAMA']['en'] + self.stats['SAMA']['ar']
        self.cma_count = self.stats['CMA']['en'] + self.stats['CMA']['ar']
        self.total = len(self.documents)
        
        print(f"    ✓ {self.total} articles (SAMA: {self.sama_count}, CMA: {self.cma_count})")
        
        # Load BM25 index
        print("  Step 3: Loading BM25...")
        with open(BM25_PATH, 'rb') as f:
            self.bm25 = pickle.load(f)
        print("    ✓ BM25 ready")
        
        # Load embedding model
        print("  Loading embedding model...")
        self.embedder = SentenceTransformer(EMBEDDING_MODEL)
        print(f"    ✓ {EMBEDDING_MODEL}")
        
        # Connect to ChromaDB
        print("  Connecting to ChromaDB...")
        self.client = chromadb.PersistentClient(path=CHROMA_PATH)
        self.collection = self.client.get_collection("tadqeeq_v2")
        print(f"    ✓ {self.collection.count()} documents")
        
        # Warm up LLM
        print("  Loading LLM...")
        self._warmup_llm()
        print(f"    ✓ {LLM_MODEL}")
        
        # Chat history
        self.chat_history = ChatHistory()

        self.doc_processor = DocumentProcessor()
        self.compliance_checker = ComplianceChecker()
        self.chat_exporter = ChatExporter()
        
        print("\n✓ TadqeeqAI v2.2 Ready!\n")
    
    def _warmup_llm(self):
        try:
            ollama.generate(model=LLM_MODEL, prompt="test", options={"num_predict": 1})
        except:
            print(f"      Pulling {LLM_MODEL}...")
            ollama.pull(LLM_MODEL)
    
    def detect_language(self, text):
        arabic_chars = len(re.findall(r'[\u0600-\u06FF]', text))
        return 'ar' if arabic_chars > len(text) * 0.3 else 'en'
    
    def detect_regulator(self, query):
        q_lower = query.lower()
        
        # English keywords
        sama_en = ['sama', 'finance company', 'finance companies', 'financing company',
                   'licensing fee', 'real estate finance', 'mortgage', 'microfinance',
                   'finance control', 'monetary authority', 'bank', 'banking']
        cma_en = ['cma', 'capital market', 'securities', 'sukuk', 'debt instrument',
                  'investment fund', 'qualified investor', 'public offering', 'ipo',
                  'private placement', 'prospectus', 'listing', 'merger', 'acquisition',
                  'stock', 'shares', 'exchange']
        
        # Arabic keywords (don't lowercase)
        sama_ar = ['ساما', 'شركة التمويل', 'شركات التمويل', 'رسوم الترخيص', 'التمويل العقاري',
                   'التمويل الأصغر', 'مؤسسة النقد', 'البنك المركزي', 'تمويل']
        cma_ar = ['مستثمر مؤهل', 'المستثمر المؤهل', 'هيئة السوق المالية', 'هيئة السوق',
                  'صكوك', 'الصكوك', 'طرح عام', 'طرح خاص', 'نشرة الإصدار',
                  'صناديق الاستثمار', 'أوراق مالية', 'الأوراق المالية', 'سوق المال',
                  'الاندماج', 'الاستحواذ', 'الأسهم', 'التداول']
        
        # Check English keywords
        sama_match = any(kw in q_lower for kw in sama_en)
        cma_match = any(kw in q_lower for kw in cma_en)
        
        # Check Arabic keywords (case-sensitive)
        sama_match = sama_match or any(kw in query for kw in sama_ar)
        cma_match = cma_match or any(kw in query for kw in cma_ar)
        
        if sama_match and cma_match:
            return 'BOTH'
        elif sama_match:
            return 'SAMA'
        elif cma_match:
            return 'CMA'
        return 'BOTH'
    
    def translate_arabic_query(self, query):
        """Translate Arabic query keywords to English for searching English documents."""
        translations = {
            'رسوم الترخيص': 'licensing fees',
            'رسوم ترخيص': 'licensing fees',
            'شركات التمويل': 'finance companies',
            'شركة التمويل': 'finance company',
            'التمويل العقاري': 'real estate finance',
            'التمويل الأصغر': 'microfinance',
            'المستثمر المؤهل': 'qualified investor',
            'مستثمر مؤهل': 'qualified investor',
            'الصكوك': 'sukuk debt instruments',
            'صكوك': 'sukuk debt instruments',
            'أدوات الدين': 'debt instruments',
            'طرح عام': 'public offering',
            'طرح خاص': 'private placement',
            'نشرة الإصدار': 'prospectus',
            'صناديق الاستثمار': 'investment funds',
            'صندوق استثمار': 'investment fund',
            'رأس المال': 'capital requirements',
            'متطلبات رأس المال': 'capital requirements',
            'الحد الأدنى': 'minimum requirements',
            'هيئة السوق المالية': 'capital market authority CMA',
            'مؤسسة النقد': 'SAMA monetary authority',
            'ساما': 'SAMA',
            'الاندماج': 'merger',
            'الاستحواذ': 'acquisition',
            'الأسهم': 'shares stocks',
            'الإفصاح': 'disclosure',
            'الحوكمة': 'governance',
            'مجلس الإدارة': 'board of directors',
            'تقرير سنوي': 'annual report',
            'القوائم المالية': 'financial statements',
            'المراجع الخارجي': 'external auditor',
            'العقوبات': 'penalties',
            'المخالفات': 'violations',
            'الترخيص': 'license licensing',
            'التسجيل': 'registration',
            'الإدراج': 'listing',
            'السوق الموازية': 'parallel market',
            'الطرح': 'offering',
            'الاكتتاب': 'subscription IPO',
        }
        
        result = query
        for ar, en in translations.items():
            if ar in query:
                result = result + ' ' + en
        
        if result != query:
            return result
        
        return query
    
    def expand_query(self, query, lang):
        q = query.lower()
        mappings = {
            'sukuk': 'debt instruments securities bonds',
            'sukuk issuance': 'debt instruments offering securities',
            'debt instruments': 'sukuk securities bonds',
            'licensing fee': 'license fee financial consideration',
            'licensing fees': 'license fee financial consideration',
            'qualified investor': 'accredited investor',
            'capital requirements': 'minimum capital paid up capital',
            'finance company': 'finance companies',
            'microfinance': 'micro finance small finance',
            'real estate finance': 'mortgage property finance',
            'investment fund': 'investment funds',
            'public offering': 'IPO offering securities',
            'private placement': 'exempt offering',
        }
        expansions = [exp for term, exp in mappings.items() if term in q]
        return query + ' ' + ' '.join(expansions) if expansions else query
    
    def bm25_search(self, query, regulator, language, top_k=15, force_english=False):
        """BM25 keyword search."""
        search_lang = 'en' if force_english else language
        tokens = re.findall(r'[\u0600-\u06FF]+|[a-zA-Z]+|\d+', query.lower())
        if not tokens: return []
        scores = self.bm25.get_scores(tokens)
        results = []
        for idx in np.argsort(scores)[::-1]:
            if len(results) >= top_k or idx >= len(self.documents): continue
            doc = self.documents[idx]
            if regulator != 'BOTH' and doc.get('regulator') != regulator: continue
            if doc.get('language') != search_lang: continue
            if scores[idx] > 0:
                results.append({'doc': doc, 'score': float(scores[idx]), 'source': 'bm25'})
        return results
    
    def semantic_search(self, query, regulator, language, top_k=15, force_english=False):
        """Semantic search."""
        search_lang = 'en' if force_english else language
        embedding = self.embedder.encode([f"query: {query}"]).tolist()
        where = {"language": {"$eq": search_lang}} if regulator == 'BOTH' else {
            "$and": [{"language": {"$eq": search_lang}}, {"regulator": {"$eq": regulator}}]}
        try:
            results = self.collection.query(query_embeddings=embedding, n_results=top_k, where=where)
        except:
            results = self.collection.query(query_embeddings=embedding, n_results=top_k * 2)
        output = []
        if results['documents'] and results['documents'][0]:
            for doc_text, meta, dist in zip(results['documents'][0], results['metadatas'][0], results['distances'][0]):
                if regulator != 'BOTH' and meta.get('regulator') != regulator: continue
                if meta.get('language') != search_lang: continue
                output.append({'doc': {'text': doc_text, 'article': meta.get('article', ''),
                    'document': meta.get('document', ''), 'regulator': meta.get('regulator', ''),
                    'language': meta.get('language', '')}, 'score': 1/(1+dist), 'source': 'semantic'})
        return output[:top_k]
    
    def hybrid_search(self, query, n_results=3):
        """Hybrid search with English Bridge strategy."""
        user_language = self.detect_language(query)
        regulator = self.detect_regulator(query)
        
        if user_language == 'ar':
            english_query = self.translate_arabic_query(query)
            expanded = self.expand_query(english_query, 'en')
            print(f"DEBUG: Arabic Query → English Bridge")
            print(f"  Original: {query[:50]}")
            print(f"  Translated: {english_query[:50]}")
            print(f"  Regulator: {regulator}")
            bm25_res = self.bm25_search(expanded, regulator, user_language, force_english=True)
            sem_res = self.semantic_search(expanded, regulator, user_language, force_english=True)
        else:
            expanded = self.expand_query(query, user_language)
            print(f"DEBUG: English Query='{query[:50]}', Reg={regulator}")
            bm25_res = self.bm25_search(expanded, regulator, user_language)
            sem_res = self.semantic_search(expanded, regulator, user_language)
        
        print(f"DEBUG: BM25={len(bm25_res)}, Semantic={len(sem_res)}")
        
        # RRF Fusion
        doc_scores = {}
        k = 60
        for rank, r in enumerate(bm25_res):
            key = f"{r['doc']['document']}:{r['doc']['article']}"
            if key not in doc_scores: doc_scores[key] = {'doc': r['doc'], 'rrf': 0, 'src': set()}
            doc_scores[key]['rrf'] += 1/(k+rank+1)
            doc_scores[key]['src'].add('BM25')
        for rank, r in enumerate(sem_res):
            key = f"{r['doc']['document']}:{r['doc']['article']}"
            if key not in doc_scores: doc_scores[key] = {'doc': r['doc'], 'rrf': 0, 'src': set()}
            doc_scores[key]['rrf'] += 1/(k+rank+1)
            doc_scores[key]['src'].add('Semantic')
        sorted_res = sorted(doc_scores.values(), key=lambda x: x['rrf'], reverse=True)
        final = []
        for r in sorted_res[:n_results]:
            final.append(r['doc'])
            print(f"  → {r['doc']['article']} [{'+'.join(r['src'])}]")
        return final, regulator, user_language
    
    def is_follow_up(self, query):
        """Detect if query is a follow-up request."""
        query_lower = query.lower().strip()
        
        # English follow-up patterns
        follow_up_en = [
            'yes', 'yeah', 'sure', 'please', 'ok', 'okay',
            'simplify', 'explain', 'example', 'examples', 'scenario',
            'more details', 'elaborate', 'clarify', 'what do you mean',
            'can you explain', 'help me understand', 'break it down',
            'in simple terms', 'simpler', 'easier'
        ]
        
        # Arabic follow-up patterns
        follow_up_ar = [
            'نعم', 'أجل', 'طيب', 'حسنا', 'موافق', 'تمام',
            'وضح', 'اشرح', 'مثال', 'أمثلة', 'سيناريو',
            'تفاصيل أكثر', 'بسط', 'بشكل أبسط', 'ساعدني أفهم',
            'ماذا تعني', 'اشرح أكثر'
        ]
        
        if any(pattern in query_lower for pattern in follow_up_en):
            return True
        if any(pattern in query for pattern in follow_up_ar):
            return True
        
        # Very short responses are likely follow-ups
        if len(query.strip()) < 15 and len(query.split()) <= 3:
            return True
        
        return False
    
    def is_out_of_domain(self, query):
        """Check if query is outside SAMA/CMA domain."""
        query_lower = query.lower()
        
        # Out of domain indicators
        out_of_domain = [
            'weather', 'recipe', 'cook', 'movie', 'song', 'music',
            'game', 'sport', 'football', 'soccer', 'basketball',
            'joke', 'story', 'poem', 'write me', 'create a',
            'translate', 'what is the capital', 'who is the president',
            'how to code', 'python', 'javascript', 'programming',
            'health', 'medical', 'doctor', 'disease',
            'travel', 'hotel', 'flight', 'vacation',
            'الطقس', 'وصفة', 'طبخ', 'فيلم', 'أغنية', 'موسيقى',
            'لعبة', 'رياضة', 'كرة القدم', 'نكتة', 'قصة', 'قصيدة',
            'ترجم', 'عاصمة', 'رئيس', 'برمجة', 'صحة', 'طبيب', 'سفر'
        ]
        
        # Check for out-of-domain content
        if any(term in query_lower for term in out_of_domain):
            return True
        
        return False
    
    def build_prompt(self, question, docs, language, is_follow_up=False, conversation_context=None):
        """Build LLM prompt optimized for Aya model."""
        ctx = "\n\n---\n\n".join([f"[Document {i}]\nSource: {d['document']}\nArticle: {d['article']}\nContent:\n{d['text']}" 
                                  for i, d in enumerate(docs, 1)])
        
        # Add conversation context for follow-ups
        conv_context = ""
        if is_follow_up and conversation_context:
            conv_context = "\n\nPrevious conversation:\n"
            for msg in conversation_context[-4:]:  # Last 4 messages
                role = "User" if msg['role'] == 'user' else "Assistant"
                conv_context += f"{role}: {msg['content'][:500]}\n"
        
        if language == 'ar':
            if is_follow_up:
                return f"""أنت مساعد قانوني متخصص في الأنظمة المالية السعودية (ساما وهيئة السوق المالية).

المحادثة السابقة:
{conv_context}

المستندات المرجعية:
{ctx}

طلب المستخدم: {question}

المستخدم يطلب توضيحاً أو تبسيطاً. قم بما يلي:
- إذا طلب تبسيط: اشرح المفهوم بلغة سهلة وواضحة
- إذا طلب مثال: قدم سيناريو عملي يوضح التطبيق
- إذا طلب توضيح: اشرح النقاط الغامضة بالتفصيل

الإجابة:"""
            else:
                return f"""أنت مساعد قانوني متخصص في الأنظمة المالية السعودية (ساما وهيئة السوق المالية).

تعليمات مهمة:
- اقرأ كل مستند بعناية قبل الإجابة
- استخرج المعلومات المطلوبة من المستندات
- اذكر رقم المادة عند الاستشهاد (مثال: المادة 22)
- مصطلح "debt instruments" يعني "الصكوك" أو "أدوات الدين"
- اكتب الأرقام والمبالغ كما هي في المستند
- استخدم تنسيق Markdown: استخدم **نص** للتأكيد و - للقوائم
- إذا لم تجد المعلومة المحددة، قل ذلك بوضوح
- لا تذكر أنك ستساعد في نهاية الإجابة

المستندات المرجعية:
{ctx}

السؤال: {question}

الإجابة:"""
        
        # English prompts
        if is_follow_up:
            return f"""You are a legal assistant specializing in Saudi Arabian financial regulations (SAMA and CMA).

Previous conversation:
{conv_context}

Reference Documents:
{ctx}

User request: {question}

The user is asking for clarification or simplification. Do the following:
- If they want simplification: Explain the concept in plain, easy-to-understand language
- If they want examples: Provide a practical scenario showing how this applies
- If they want clarification: Explain the unclear points in detail

Answer:"""
        else:
            return f"""You are a legal assistant specializing in Saudi Arabian financial regulations (SAMA and CMA).

Important instructions:
- Read each document carefully before answering
- Extract the relevant information from the documents
- Cite the Article number when referencing information (e.g., "Article 22")
- "Sukuk" and "debt instruments" refer to the same thing
- Preserve exact numbers and amounts as written in the documents
- Use Markdown formatting: **bold** for emphasis and - for lists
- If the specific information is not found, say so clearly
- Do not offer further assistance at the end of your response

Reference Documents:
{ctx}

Question: {question}

Answer:"""
    
    def build_out_of_domain_response(self, language):
        """Return response for out-of-domain queries."""
        if language == 'ar':
            return """عذراً، أنا مساعد متخصص في الأنظمة المالية السعودية فقط.

يمكنني مساعدتك في:
- **أنظمة ساما**: شركات التمويل، التمويل العقاري، التمويل الأصغر
- **أنظمة هيئة السوق المالية**: الصكوك، صناديق الاستثمار، المستثمر المؤهل، الطرح والإدراج

يرجى طرح سؤال يتعلق بهذه المواضيع."""
        else:
            return """I apologize, but I am a specialized assistant for Saudi Arabian financial regulations only.

I can help you with:
- **SAMA regulations**: Finance companies, real estate finance, microfinance
- **CMA regulations**: Sukuk, investment funds, qualified investors, offerings and listings

Please ask a question related to these topics."""
    
    def generate_response(self, question):
        """Generate response with follow-up support."""
        lang = self.detect_language(question)
        
        # Check if out of domain
        if self.is_out_of_domain(question):
            return {
                'answer': self.build_out_of_domain_response(lang),
                'sources': [],
                'regulator': 'NONE'
            }
        
        # Check if this is a follow-up
        is_followup = self.is_follow_up(question)
        conversation_context = None
        
        if is_followup:
            conversation_context = self.chat_history.get_conversation_context()
            # If we have context, use the previous search results
            if conversation_context:
                # Find the last assistant message with sources
                last_sources = None
                for msg in reversed(conversation_context):
                    if msg['role'] == 'assistant' and 'sources' in msg:
                        last_sources = msg.get('sources', [])
                        break
        
        # Perform search
        docs, reg, lang = self.hybrid_search(question)
        
        if not docs:
            no_info = 'No relevant information found.' if lang == 'en' else 'لم يتم العثور على معلومات ذات صلة.'
            return {'answer': no_info, 'sources': [], 'regulator': reg}
        
        prompt = self.build_prompt(question, docs, lang, is_followup, conversation_context)
        resp = ollama.generate(model=LLM_MODEL, prompt=prompt, options={'temperature': 0.1, 'num_predict': 2000})
        
        seen = set()
        sources = [{'article': d['article'], 'document': d['document']} for d in docs if d['article'] not in seen and not seen.add(d['article'])]
        
        return {'answer': resp['response'].strip(), 'sources': sources, 'regulator': reg}


class API:
    def __init__(self):
        self.rag = None
        self.window = None  # Will be set after window creation
    
    def set_window(self, window):
        self.window = window
    
    def upload_document(self, file_data, filename):
        if not self.rag:
            self.rag = TadqeeqRAG.get_instance()
        return self.rag.doc_processor.process_file(file_data=file_data, filename=filename)
    
    def run_compliance_check(self):
        if not self.rag:
            self.rag = TadqeeqRAG.get_instance()
        text = self.rag.doc_processor.get_current_text()
        if not text:
            return {'error': 'No document uploaded'}
        return self.rag.compliance_checker.check_compliance(text, self.rag.doc_processor.current_filename)
    
    def clear_document(self):
        if not self.rag:
            self.rag = TadqeeqRAG.get_instance()
        self.rag.doc_processor.clear()
        return {'success': True}
    
    def export_markdown(self):
        print("API.export_markdown() called")
        try:
            if not self.rag:
                self.rag = TadqeeqRAG.get_instance()
            messages = self.rag.chat_history.get_messages()
            print(f"  Messages count: {len(messages) if messages else 0}")
            if not messages:
                return {'error': 'No messages to export. Start a conversation first.'}
            md_content, error = self.rag.chat_exporter.export_markdown(messages)
            print(f"  Export result - content length: {len(md_content) if md_content else 0}, error: {error}")
            if error:
                return {'error': error}
            
            # Use file dialog to save
            filename = f'tadqeeq_chat_{datetime.now().strftime("%Y%m%d_%H%M")}.md'
            if self.window:
                try:
                    save_path = self.window.create_file_dialog(
                        webview.SAVE_DIALOG,
                        save_filename=filename,
                        file_types=('Markdown Files (*.md)',)
                    )
                except:
                    # Fallback for newer pywebview versions
                    save_path = self.window.create_file_dialog(
                        dialog_type=webview.SAVE_DIALOG,
                        save_filename=filename,
                        file_types=('Markdown Files (*.md)',)
                    )
                if save_path:
                    # save_path is a tuple, get first element
                    path = save_path if isinstance(save_path, str) else save_path[0] if save_path else None
                    if path:
                        with open(path, 'w', encoding='utf-8') as f:
                            f.write(md_content)
                        return {'success': True, 'path': path}
                return {'error': 'Export cancelled'}
            else:
                # Fallback: save to current directory
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(md_content)
                return {'success': True, 'path': filename}
        except Exception as e:
            print(f"  Export error: {e}")
            import traceback
            traceback.print_exc()
            return {'error': str(e)}
    
    def export_pdf(self):
        print("API.export_pdf() called")
        try:
            if not self.rag:
                self.rag = TadqeeqRAG.get_instance()
            messages = self.rag.chat_history.get_messages()
            print(f"  Messages count: {len(messages) if messages else 0}")
            if not messages:
                return {'error': 'No messages to export. Start a conversation first.'}
            pdf_data, error = self.rag.chat_exporter.export_pdf(messages)
            print(f"  Export result - content length: {len(pdf_data) if pdf_data else 0}, error: {error}")
            if error:
                return {'error': error}
            
            # Decode base64 PDF data
            pdf_bytes = base64.b64decode(pdf_data)
            
            # Use file dialog to save
            filename = f'tadqeeq_chat_{datetime.now().strftime("%Y%m%d_%H%M")}.pdf'
            if self.window:
                try:
                    save_path = self.window.create_file_dialog(
                        webview.SAVE_DIALOG,
                        save_filename=filename,
                        file_types=('PDF Files (*.pdf)',)
                    )
                except:
                    save_path = self.window.create_file_dialog(
                        dialog_type=webview.SAVE_DIALOG,
                        save_filename=filename,
                        file_types=('PDF Files (*.pdf)',)
                    )
                if save_path:
                    path = save_path if isinstance(save_path, str) else save_path[0] if save_path else None
                    if path:
                        with open(path, 'wb') as f:
                            f.write(pdf_bytes)
                        return {'success': True, 'path': path}
                return {'error': 'Export cancelled'}
            else:
                # Fallback: save to current directory
                with open(filename, 'wb') as f:
                    f.write(pdf_bytes)
                return {'success': True, 'path': filename}
        except Exception as e:
            print(f"  Export error: {e}")
            import traceback
            traceback.print_exc()
            return {'error': str(e)}
    
    def initialize(self):
        try:
            print("1. API.initialize() called...")
            self.rag = TadqeeqRAG.get_instance()
            print("2. RAG instance created successfully")
            return {
                'status': 'ready',
                'total': self.rag.total,
                'sama': self.rag.sama_count,
                'cma': self.rag.cma_count,
                'sama_en': self.rag.stats['SAMA']['en'],
                'sama_ar': self.rag.stats['SAMA']['ar'],
                'cma_en': self.rag.stats['CMA']['en'],
                'cma_ar': self.rag.stats['CMA']['ar'],
                'chats': self.rag.chat_history.get_recent_chats()
            }
        except Exception as e:
            print(f"API.initialize() ERROR: {e}")
            import traceback
            traceback.print_exc()
            return {'status': 'error', 'message': str(e)}
    
    def query(self, question):
        if not self.rag:
            self.rag = TadqeeqRAG.get_instance()
        
        # Add user message to history
        self.rag.chat_history.add_message('user', question)
        
        # Generate response
        result = self.rag.generate_response(question)
        
        # Add assistant message to history
        self.rag.chat_history.add_message(
            'assistant',
            result['answer'],
            result.get('sources'),
            result.get('regulator')
        )
        
        return result
    
    def new_chat(self):
        if not self.rag:
            self.rag = TadqeeqRAG.get_instance()
        # Clear any uploaded document when starting new chat
        self.rag.doc_processor.clear()
        chat_id = self.rag.chat_history.new_chat()
        return {'id': chat_id, 'chats': self.rag.chat_history.get_recent_chats()}
    
    def load_chat(self, chat_id):
        if not self.rag:
            self.rag = TadqeeqRAG.get_instance()
        messages = self.rag.chat_history.load_chat(chat_id)
        return {'messages': messages}
    
    def get_chats(self):
        if not self.rag:
            self.rag = TadqeeqRAG.get_instance()
        return {'chats': self.rag.chat_history.get_recent_chats()}
    
    def delete_chat(self, chat_id):
        if not self.rag:
            self.rag = TadqeeqRAG.get_instance()
        success = self.rag.chat_history.delete_chat(chat_id)
        return {'success': success, 'chats': self.rag.chat_history.get_recent_chats()}


HTML = '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>TadqeeqAI v2.1</title>
    <link href="https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700&display=swap" rel="stylesheet">
    <script src="https://cdn.jsdelivr.net/npm/marked/marked.min.js"></script>
    <style>
        :root {
            --bg: #0a0d10;
            --bg2: #0f1318;
            --bg3: #161b22;
            --bg4: #1c2128;
            --border: #30363d;
            --text: #e6edf3;
            --text2: #8b949e;
            --text3: #6e7681;
            --accent: #00d4aa;
            --accent2: #00b894;
            --sama: #58a6ff;
            --cma: #3fb950;
            --danger: #f85149;
        }
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: 'Plus Jakarta Sans', -apple-system, BlinkMacSystemFont, sans-serif;
            background: var(--bg);
            color: var(--text);
            height: 100vh;
            display: flex;
            -webkit-user-select: none;
            user-select: none;
        }
        
        /* Custom Scrollbar */
        ::-webkit-scrollbar { width: 8px; height: 8px; }
        ::-webkit-scrollbar-track { background: transparent; }
        ::-webkit-scrollbar-thumb { background: var(--border); border-radius: 4px; }
        ::-webkit-scrollbar-thumb:hover { background: var(--text3); }
        
        ::selection { background: var(--accent); color: var(--bg); }
        
        /* V2.2: Header Actions */
        .header-actions { display: flex; align-items: center; gap: 12px; }
        .header-btn {
            display: flex; align-items: center; gap: 6px; padding: 6px 12px;
            background: var(--bg3); border: 1px solid var(--border); border-radius: 6px;
            color: var(--text2); font-size: 12px; cursor: pointer; transition: all 0.15s ease;
        }
        .header-btn:hover { background: var(--bg4); color: var(--text); border-color: var(--accent); }
        .header-btn svg { width: 14px; height: 14px; stroke: currentColor; fill: none; }
        
        /* V2.2: Export Menu */
        .export-menu {
            position: absolute; right: 0; top: 100%; margin-top: 4px;
            background: var(--bg3); border: 1px solid var(--border); border-radius: 8px;
            padding: 4px; z-index: 100; display: none; min-width: 160px;
            box-shadow: 0 8px 24px rgba(0,0,0,0.4);
        }
        .export-menu.show { display: block; }
        .export-item {
            display: flex; align-items: center; gap: 8px; padding: 8px 12px;
            font-size: 12px; color: var(--text); border-radius: 4px;
            cursor: pointer; transition: all 0.15s ease;
        }
        .export-item:hover { background: var(--bg4); }
        .export-item svg { width: 14px; height: 14px; stroke: currentColor; fill: none; }
        
        /* V2.2: Drop Zone Overlay */
        .drop-overlay {
            position: absolute; inset: 0; background: rgba(0, 212, 170, 0.1);
            border: 2px dashed var(--accent); display: none; align-items: center;
            justify-content: center; z-index: 50; pointer-events: none;
        }
        .drop-overlay.active { display: flex; pointer-events: auto; }
        .drop-overlay-content {
            background: var(--bg2); padding: 40px 60px; border-radius: 16px;
            text-align: center; border: 1px solid var(--accent); pointer-events: none;
        }
        .drop-overlay-content svg { width: 48px; height: 48px; stroke: var(--accent); margin-bottom: 16px; }
        .drop-overlay-content h3 { color: var(--text); margin-bottom: 8px; }
        .drop-overlay-content p { color: var(--text2); font-size: 13px; }
        
        /* V2.2: Document Badge */
        .doc-badge {
            display: flex; align-items: center; gap: 10px;
            background: var(--bg3); padding: 10px 14px; border-radius: 8px;
            border: 1px solid var(--border);
        }
        .doc-badge svg { width: 18px; height: 18px; stroke: var(--accent); fill: none; flex-shrink: 0; }
        .doc-badge-info { flex: 1; min-width: 0; }
        .doc-badge-name { font-size: 13px; color: var(--text); font-weight: 500; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }
        .doc-badge-meta { font-size: 11px; color: var(--text3); margin-top: 2px; }
        .doc-badge-actions { display: flex; gap: 6px; }
        .doc-badge-btn {
            padding: 5px 10px; border-radius: 5px; font-size: 11px;
            cursor: pointer; border: 1px solid var(--border); transition: all 0.15s ease;
        }
        .doc-badge-btn.check { background: var(--accent); color: var(--bg); border-color: var(--accent); font-weight: 500; }
        .doc-badge-btn.check:hover { opacity: 0.9; }
        .doc-badge-btn.clear { background: transparent; color: var(--text2); }
        .doc-badge-btn.clear:hover { background: var(--bg4); color: var(--danger); }
        
        /* V2.2: Compliance Report */
        .compliance-report {
            background: var(--bg2); border: 1px solid var(--border); border-radius: 12px;
            padding: 16px; margin: 12px 0; max-width: 600px;
        }
        .compliance-header { display: flex; justify-content: space-between; align-items: center; margin-bottom: 16px; flex-wrap: wrap; gap: 8px; }
        .compliance-title { font-size: 14px; font-weight: 600; color: var(--text); }
        .compliance-score { padding: 6px 12px; border-radius: 20px; font-size: 12px; font-weight: 600; }
        .compliance-score.good { background: rgba(63, 185, 80, 0.15); color: var(--cma); }
        .compliance-score.warning { background: rgba(210, 153, 34, 0.15); color: #d29922; }
        .compliance-score.bad { background: rgba(248, 81, 73, 0.15); color: var(--danger); }
        .compliance-item { display: flex; align-items: flex-start; gap: 12px; padding: 10px 0; border-bottom: 1px solid var(--border); }
        .compliance-item:last-child { border-bottom: none; }
        .compliance-icon { width: 22px; height: 22px; border-radius: 50%; display: flex; align-items: center; justify-content: center; flex-shrink: 0; }
        .compliance-icon.compliant { background: rgba(63, 185, 80, 0.2); color: var(--cma); }
        .compliance-icon.warning { background: rgba(210, 153, 34, 0.2); color: #d29922; }
        .compliance-icon svg { width: 12px; height: 12px; stroke: currentColor; fill: none; }
        .compliance-content { flex: 1; }
        .compliance-name { font-size: 13px; font-weight: 500; color: var(--text); }
        .compliance-reg { font-size: 11px; color: var(--text3); margin-top: 2px; }
        .compliance-detail { font-size: 12px; color: var(--text2); margin-top: 4px; }
        
        /* Sidebar */
        .sidebar {
            width: 280px;
            background: var(--bg2);
            border-right: 1px solid var(--border);
            display: flex;
            flex-direction: column;
        }
        .sidebar-header {
            padding: 16px;
            border-bottom: 1px solid var(--border);
        }
        .logo {
            display: flex;
            align-items: center;
            gap: 12px;
            margin-bottom: 12px;
        }
        .logo-icon {
            width: 38px;
            height: 38px;
            background: linear-gradient(135deg, var(--accent), var(--accent2));
            border-radius: 10px;
            display: flex;
            align-items: center;
            justify-content: center;
        }
        .logo-icon svg { width: 20px; height: 20px; fill: var(--bg); }
        .logo-text { font-size: 17px; font-weight: 700; letter-spacing: -0.3px; }
        .logo-sub { font-size: 10px; color: var(--text3); margin-top: 2px; letter-spacing: 0.5px; }
        
        /* New Chat Button */
        .new-chat-btn {
            width: 100%;
            padding: 10px 14px;
            background: var(--bg3);
            border: 1px solid var(--border);
            border-radius: 8px;
            color: var(--text);
            font-size: 13px;
            font-weight: 500;
            cursor: pointer;
            display: flex;
            align-items: center;
            gap: 8px;
            transition: all 0.15s ease;
        }
        .new-chat-btn:hover {
            background: var(--bg4);
            border-color: var(--accent);
        }
        .new-chat-btn svg { width: 16px; height: 16px; stroke: var(--accent); fill: none; }
        
        /* Chat History */
        .chat-history {
            flex: 1;
            overflow-y: auto;
            padding: 12px;
        }
        .history-title {
            font-size: 10px;
            text-transform: uppercase;
            letter-spacing: 1.5px;
            color: var(--text3);
            margin-bottom: 8px;
            font-weight: 600;
            padding: 0 4px;
        }
        .history-item {
            display: flex;
            align-items: center;
            padding: 8px 10px;
            background: transparent;
            border-radius: 8px;
            margin-bottom: 4px;
            font-size: 12px;
            color: var(--text2);
            cursor: pointer;
            transition: all 0.15s ease;
            border: 1px solid transparent;
            position: relative;
        }
        .history-item:hover {
            background: var(--bg3);
            color: var(--text);
        }
        .history-item.active {
            background: var(--bg4);
            border-color: var(--border);
            color: var(--text);
        }
        .history-item-text {
            flex: 1;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
        }
        .history-item-menu {
            opacity: 0;
            width: 24px;
            height: 24px;
            border-radius: 4px;
            display: flex;
            align-items: center;
            justify-content: center;
            cursor: pointer;
            transition: all 0.15s ease;
            flex-shrink: 0;
        }
        .history-item:hover .history-item-menu { opacity: 1; }
        .history-item-menu:hover { background: var(--bg4); }
        .history-item-menu svg { width: 14px; height: 14px; fill: var(--text3); }
        
        /* Dropdown Menu */
        .dropdown {
            position: absolute;
            right: 0;
            top: 100%;
            background: var(--bg3);
            border: 1px solid var(--border);
            border-radius: 8px;
            padding: 4px;
            z-index: 100;
            display: none;
            min-width: 120px;
            box-shadow: 0 8px 24px rgba(0,0,0,0.4);
        }
        .dropdown.show { display: block; }
        .dropdown-item {
            display: flex;
            align-items: center;
            gap: 8px;
            padding: 8px 12px;
            font-size: 12px;
            color: var(--danger);
            border-radius: 4px;
            cursor: pointer;
            transition: all 0.15s ease;
        }
        .dropdown-item:hover { background: var(--bg4); }
        .dropdown-item svg { width: 14px; height: 14px; stroke: currentColor; fill: none; }
        
        /* Delete Confirmation Modal */
        .modal-overlay {
            position: fixed;
            inset: 0;
            background: rgba(0,0,0,0.7);
            display: none;
            align-items: center;
            justify-content: center;
            z-index: 200;
        }
        .modal-overlay.show { display: flex; }
        .modal {
            background: var(--bg2);
            border: 1px solid var(--border);
            border-radius: 12px;
            padding: 24px;
            max-width: 400px;
            width: 90%;
            box-shadow: 0 16px 48px rgba(0,0,0,0.5);
        }
        .modal-title {
            font-size: 16px;
            font-weight: 600;
            margin-bottom: 12px;
            color: var(--text);
        }
        .modal-text {
            font-size: 13px;
            color: var(--text2);
            margin-bottom: 20px;
            line-height: 1.5;
        }
        .modal-buttons {
            display: flex;
            gap: 10px;
            justify-content: flex-end;
        }
        .modal-btn {
            padding: 8px 16px;
            border-radius: 6px;
            font-size: 13px;
            font-weight: 500;
            cursor: pointer;
            transition: all 0.15s ease;
            border: 1px solid var(--border);
        }
        .modal-btn.cancel {
            background: var(--bg3);
            color: var(--text);
        }
        .modal-btn.cancel:hover { background: var(--bg4); }
        .modal-btn.delete {
            background: var(--danger);
            color: white;
            border-color: var(--danger);
        }
        .modal-btn.delete:hover { opacity: 0.9; }
        
        /* Examples Section */
        .examples { padding: 12px; border-top: 1px solid var(--border); }
        .ex-title {
            font-size: 10px;
            text-transform: uppercase;
            letter-spacing: 1.5px;
            color: var(--text3);
            margin-bottom: 8px;
            font-weight: 600;
        }
        .ex {
            padding: 8px 10px;
            background: var(--bg3);
            border-radius: 6px;
            margin-bottom: 4px;
            font-size: 11px;
            color: var(--text2);
            cursor: pointer;
            transition: all 0.15s ease;
            border: 1px solid transparent;
            display: flex;
            justify-content: space-between;
            align-items: center;
        }
        .ex:hover {
            background: var(--bg4);
            border-color: var(--border);
            color: var(--text);
        }
        .tag {
            font-size: 8px;
            padding: 2px 5px;
            border-radius: 3px;
            font-weight: 600;
            text-transform: uppercase;
            flex-shrink: 0;
        }
        .tag.sama { background: rgba(88, 166, 255, 0.15); color: var(--sama); }
        .tag.cma { background: rgba(63, 185, 80, 0.15); color: var(--cma); }
        .lang-section { margin-top: 10px; }
        .lang-section .ex { direction: rtl; text-align: right; }
        
        /* Status */
        .status {
            padding: 12px 14px;
            border-top: 1px solid var(--border);
            font-size: 10px;
            color: var(--text3);
            background: var(--bg);
        }
        .status-row { display: flex; align-items: center; gap: 8px; }
        .dot {
            width: 6px;
            height: 6px;
            background: var(--accent);
            border-radius: 50%;
            box-shadow: 0 0 8px var(--accent);
        }
        .dot.loading { animation: pulse 1.5s infinite; }
        @keyframes pulse { 50% { opacity: 0.3; } }
        
        /* Main Area */
        .main {
            flex: 1;
            display: flex;
            flex-direction: column;
            min-width: 0;
            background: var(--bg);
        }
        .header {
            padding: 10px 20px;
            border-bottom: 1px solid var(--border);
            display: flex;
            justify-content: space-between;
            align-items: center;
            background: var(--bg2);
        }
        .header-title { font-size: 12px; color: var(--text2); font-weight: 500; }
        .badge {
            padding: 4px 10px;
            background: var(--bg3);
            border-radius: 12px;
            font-size: 10px;
            color: var(--text3);
            font-weight: 500;
            border: 1px solid var(--border);
        }
        
        /* Chat Area */
        .chat {
            flex: 1;
            overflow-y: auto;
            padding: 20px;
            -webkit-user-select: text;
            user-select: text;
        }
        
        /* Welcome Screen */
        .welcome {
            display: flex;
            flex-direction: column;
            align-items: center;
            justify-content: center;
            height: 100%;
            text-align: center;
        }
        .welcome-icon {
            width: 64px;
            height: 64px;
            background: linear-gradient(135deg, var(--accent), var(--accent2));
            border-radius: 18px;
            display: flex;
            align-items: center;
            justify-content: center;
            margin-bottom: 16px;
            box-shadow: 0 8px 32px rgba(0, 212, 170, 0.25);
        }
        .welcome-icon svg { width: 32px; height: 32px; fill: var(--bg); }
        .welcome h1 {
            font-size: 26px;
            margin-bottom: 8px;
            background: linear-gradient(90deg, var(--accent), #7c72ff);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            font-weight: 700;
        }
        .welcome p { color: var(--text2); max-width: 440px; line-height: 1.6; font-size: 13px; }
        .welcome-stats { display: flex; gap: 10px; margin-top: 24px; }
        .stat {
            background: var(--bg3);
            padding: 12px 18px;
            border-radius: 10px;
            text-align: center;
            border: 1px solid var(--border);
            min-width: 80px;
        }
        .stat-val { font-size: 22px; font-weight: 700; color: var(--accent); }
        .stat-lbl { font-size: 9px; color: var(--text3); margin-top: 2px; text-transform: uppercase; letter-spacing: 0.5px; }
        
        /* Messages */
        .msg {
            display: flex;
            gap: 12px;
            margin-bottom: 18px;
            max-width: 820px;
            animation: fadeIn 0.25s ease;
        }
        .msg.user {
            flex-direction: row-reverse;
            margin-left: auto;
            margin-right: 0;
        }
        .msg.assistant {
            margin-left: 0;
            margin-right: auto;
        }
        @keyframes fadeIn { from { opacity: 0; transform: translateY(6px); } to { opacity: 1; transform: translateY(0); } }
        
        .avatar {
            width: 30px;
            height: 30px;
            border-radius: 8px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-weight: 600;
            font-size: 10px;
            flex-shrink: 0;
        }
        .msg.user .avatar { background: var(--bg3); color: var(--text2); border: 1px solid var(--border); }
        .msg.assistant .avatar { background: linear-gradient(135deg, var(--accent), var(--accent2)); color: var(--bg); }
        .msg.assistant .avatar svg { width: 14px; height: 14px; fill: var(--bg); }
        
        .msg-body { flex: 1; min-width: 0; }
        .msg.user .msg-body { text-align: right; }
        
        .msg-header { display: flex; align-items: center; gap: 8px; margin-bottom: 4px; }
        .msg.user .msg-header { justify-content: flex-end; }
        .msg-role { font-size: 11px; font-weight: 600; color: var(--text2); }
        
        .reg { font-size: 9px; padding: 2px 6px; border-radius: 4px; font-weight: 600; text-transform: uppercase; }
        .reg.sama { background: rgba(88, 166, 255, 0.15); color: var(--sama); }
        .reg.cma { background: rgba(63, 185, 80, 0.15); color: var(--cma); }
        .reg.both { background: rgba(168, 85, 247, 0.15); color: #a371f7; }
        
        .msg-text {
            line-height: 1.65;
            font-size: 14px;
            color: var(--text);
            -webkit-user-select: text;
            user-select: text;
        }
        .msg.user .msg-text {
            background: var(--bg3);
            padding: 10px 14px;
            border-radius: 12px 12px 4px 12px;
            display: inline-block;
            border: 1px solid var(--border);
        }
        .msg-text.rtl { direction: rtl; text-align: right; }
        .msg-text p { margin-bottom: 10px; }
        .msg-text p:last-child { margin-bottom: 0; }
        .msg-text ul, .msg-text ol { margin: 10px 0; padding-left: 20px; }
        .msg-text.rtl ul, .msg-text.rtl ol { padding-left: 0; padding-right: 20px; }
        .msg-text li { margin-bottom: 5px; }
        .msg-text strong { color: var(--accent); font-weight: 600; }
        .msg-text code { background: var(--bg3); padding: 2px 5px; border-radius: 4px; font-size: 13px; }
        .msg-text h1, .msg-text h2, .msg-text h3 { margin: 14px 0 6px; color: var(--text); font-size: 15px; }
        
        /* Sources */
        .sources { margin-top: 12px; padding-top: 10px; border-top: 1px solid var(--border); }
        .src-title { font-size: 9px; text-transform: uppercase; letter-spacing: 1px; color: var(--text3); margin-bottom: 6px; font-weight: 600; }
        .src {
            display: inline-block;
            background: var(--bg3);
            padding: 4px 8px;
            border-radius: 4px;
            font-size: 10px;
            color: var(--text2);
            margin-right: 4px;
            margin-bottom: 4px;
            border: 1px solid var(--border);
        }
        
        /* Loading */
        .loading-msg { display: flex; align-items: center; gap: 10px; color: var(--text3); font-size: 13px; }
        .dots { display: flex; gap: 3px; }
        .dots span { width: 5px; height: 5px; background: var(--accent); border-radius: 50%; animation: bounce 1.4s infinite; }
        .dots span:nth-child(2) { animation-delay: 0.2s; }
        .dots span:nth-child(3) { animation-delay: 0.4s; }
        @keyframes bounce { 0%, 80%, 100% { transform: scale(0.7); opacity: 0.4; } 40% { transform: scale(1); opacity: 1; } }
        
        /* Input Area - Gemini Style */
        .input-area {
            padding: 14px 20px 12px;
            border-top: 1px solid var(--border);
            background: var(--bg2);
        }
        .input-box {
            max-width: 820px;
            margin: 0 auto;
            display: flex;
            align-items: flex-end;
            gap: 12px;
            background: var(--bg3);
            border: 1px solid var(--border);
            border-radius: 24px;
            padding: 8px 8px 8px 16px;
            transition: all 0.2s ease;
        }
        .input-box:focus-within {
            border-color: var(--accent);
            box-shadow: 0 0 0 3px rgba(0, 212, 170, 0.08);
        }
        .input-box .attach-btn {
            width: 36px;
            height: 36px;
            border-radius: 50%;
            border: none;
            background: transparent;
            color: var(--text2);
            cursor: pointer;
            display: flex;
            align-items: center;
            justify-content: center;
            transition: all 0.15s ease;
            flex-shrink: 0;
        }
        .input-box .attach-btn:hover { background: var(--bg4); color: var(--accent); }
        .input-box .attach-btn svg { width: 20px; height: 20px; stroke: currentColor; fill: none; }
        .input-box textarea {
            flex: 1;
            background: transparent;
            border: none;
            color: var(--text);
            font-family: inherit;
            font-size: 14px;
            resize: none;
            outline: none;
            max-height: 120px;
            line-height: 1.5;
            padding: 8px 0;
            min-width: 0;
        }
        .input-box textarea::placeholder { color: var(--text3); }
        .input-box .send {
            background: linear-gradient(135deg, var(--accent), var(--accent2));
            border: none;
            border-radius: 50%;
            width: 36px;
            height: 36px;
            cursor: pointer;
            display: flex;
            align-items: center;
            justify-content: center;
            flex-shrink: 0;
            transition: all 0.15s ease;
        }
        .input-box .send:hover { transform: scale(1.05); box-shadow: 0 4px 12px rgba(0, 212, 170, 0.3); }
        .input-box .send:disabled { opacity: 0.4; cursor: not-allowed; transform: none; }
        .input-box .send svg { width: 16px; height: 16px; fill: var(--bg); }
        .input-disclaimer {
            text-align: center;
            font-size: 11px;
            color: var(--text3);
            margin-top: 8px;
            opacity: 0.7;
        }
        
        /* Overlay */
        .overlay {
            position: fixed;
            inset: 0;
            background: rgba(10, 13, 16, 0.98);
            display: flex;
            flex-direction: column;
            align-items: center;
            justify-content: center;
            gap: 16px;
            z-index: 100;
        }
        .overlay.hidden { display: none; }
        .spinner {
            width: 40px;
            height: 40px;
            border: 2px solid var(--border);
            border-top-color: var(--accent);
            border-radius: 50%;
            animation: spin 0.8s linear infinite;
        }
        @keyframes spin { to { transform: rotate(360deg); } }
        .overlay-text { color: var(--text2); font-size: 13px; }
    </style>
</head>
<body>
    <aside class="sidebar">
        <div class="sidebar-header">
            <div class="logo">
                <div class="logo-icon">
                    <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2">
                        <path d="M12 2L2 7l10 5 10-5-10-5zM2 17l10 5 10-5M2 12l10 5 10-5" stroke="currentColor" fill="none"/>
                    </svg>
                </div>
                <div>
                    <div class="logo-text">TadqeeqAI</div>
                    <div class="logo-sub">SAMA · CMA · EN/AR</div>
                </div>
            </div>
            <button class="new-chat-btn" id="newChatBtn">
                <svg viewBox="0 0 24 24" stroke-width="2"><line x1="12" y1="5" x2="12" y2="19"/><line x1="5" y1="12" x2="19" y2="12"/></svg>
                New Chat
            </button>
        </div>
        <div class="chat-history" id="chatHistory">
            <div class="history-title">Recent Chats</div>
        </div>
        <div class="examples">
            <div class="ex-title">Try These</div>
            <div class="ex" data-q="What are the licensing fees for finance companies?"><span>Licensing fees</span><span class="tag sama">SAMA</span></div>
            <div class="ex" data-q="What is a qualified investor?"><span>Qualified investor</span><span class="tag cma">CMA</span></div>
            <div class="lang-section">
                <div class="ex" data-q="ما هي رسوم ترخيص شركات التمويل؟"><span class="tag sama">SAMA</span><span>رسوم الترخيص</span></div>
                <div class="ex" data-q="ما هو المستثمر المؤهل؟"><span class="tag cma">CMA</span><span>المستثمر المؤهل</span></div>
            </div>
        </div>
        <div class="status">
            <div class="status-row"><div class="dot loading" id="dot"></div><span id="status">Initializing...</span></div>
        </div>
    </aside>
    <main class="main" style="position: relative;">
        <header class="header">
            <span class="header-title">Saudi Financial Regulations Assistant</span>
            <div class="header-actions">
                <div style="position: relative;">
                    <button class="header-btn" id="exportBtn">
                        <svg viewBox="0 0 24 24" stroke-width="2"><path d="M21 15v4a2 2 0 01-2 2H5a2 2 0 01-2-2v-4M7 10l5 5 5-5M12 15V3"/></svg>
                        Export
                    </button>
                    <div class="export-menu" id="exportMenu">
                        <div class="export-item" id="exportMd">
                            <svg viewBox="0 0 24 24" stroke-width="2"><path d="M14 2H6a2 2 0 00-2 2v16a2 2 0 002 2h12a2 2 0 002-2V8z"/><polyline points="14 2 14 8 20 8"/></svg>
                            Markdown (.md)
                        </div>
                        <div class="export-item" id="exportPdf">
                            <svg viewBox="0 0 24 24" stroke-width="2"><path d="M14 2H6a2 2 0 00-2 2v16a2 2 0 002 2h12a2 2 0 002-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="16" y1="13" x2="8" y2="13"/><line x1="16" y1="17" x2="8" y2="17"/></svg>
                            PDF (.pdf)
                        </div>
                    </div>
                </div>
                <div class="badge">v2.2 · Hybrid Search · Aya 8B</div>
            </div>
        </header>
        <!-- V2.2: Drop Zone Overlay -->
        <div class="drop-overlay" id="dropOverlay">
            <div class="drop-overlay-content">
                <svg viewBox="0 0 24 24" stroke-width="2" fill="none"><path d="M21 15v4a2 2 0 01-2 2H5a2 2 0 01-2-2v-4"/><polyline points="17 8 12 3 7 8"/><line x1="12" y1="3" x2="12" y2="15"/></svg>
                <h3>Drop your document here</h3>
                <p>PDF or DOCX (max 50 pages)</p>
            </div>
        </div>
        <div class="chat" id="chat">
            <div class="welcome" id="welcome">
                <div class="welcome-icon">
                    <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2">
                        <path d="M12 2L2 7l10 5 10-5-10-5zM2 17l10 5 10-5M2 12l10 5 10-5" stroke="currentColor" fill="none"/>
                    </svg>
                </div>
                <h1>TadqeeqAI</h1>
                <p>Bilingual AI assistant for Saudi Arabian financial regulations. Query SAMA and CMA documents in English or Arabic with accurate, cited responses.</p>
                <div class="welcome-stats">
                    <div class="stat"><div class="stat-val" id="s-sama">-</div><div class="stat-lbl">SAMA</div></div>
                    <div class="stat"><div class="stat-val" id="s-cma">-</div><div class="stat-lbl">CMA</div></div>
                    <div class="stat"><div class="stat-val" id="s-total">-</div><div class="stat-lbl">Articles</div></div>
                </div>
            </div>
        </div>
        <div class="input-area">
            <div id="docBadge" style="max-width: 820px; margin: 0 auto 10px; display: none;"></div>
            <div class="input-box">
                <button class="attach-btn" id="attachBtn" title="Upload document">
                    <svg viewBox="0 0 24 24" stroke-width="2"><path d="M21.44 11.05l-9.19 9.19a6 6 0 01-8.49-8.49l9.19-9.19a4 4 0 015.66 5.66l-9.2 9.19a2 2 0 01-2.83-2.83l8.49-8.48"/></svg>
                </button>
                <textarea id="input" placeholder="Ask about SAMA or CMA regulations..." rows="1" disabled></textarea>
                <button class="send" id="send" disabled><svg viewBox="0 0 24 24"><path d="M2 21L23 12 2 3 2 10 17 12 2 14z"/></svg></button>
            </div>
            <div class="input-disclaimer">AI can make mistakes. Please verify important information.</div>
        </div>
        <input type="file" id="fileInput" accept=".pdf,.docx,.doc" style="display: none;">
    </main>
    <div class="overlay" id="overlay"><div class="spinner"></div><div class="overlay-text">Loading TadqeeqAI...</div></div>
    
    <!-- Delete Confirmation Modal -->
    <div class="modal-overlay" id="deleteModal">
        <div class="modal">
            <div class="modal-title">Delete Chat</div>
            <div class="modal-text">Are you sure you want to delete this chat? This action is permanent and cannot be undone.</div>
            <div class="modal-buttons">
                <button class="modal-btn cancel" id="cancelDelete">Cancel</button>
                <button class="modal-btn delete" id="confirmDelete">Delete</button>
            </div>
        </div>
    </div>
    
    <!-- Error Modal -->
    <div class="modal-overlay" id="errorModal">
        <div class="modal">
            <div class="modal-title" id="errorTitle">Error</div>
            <div class="modal-text" id="errorText"></div>
            <div class="modal-buttons">
                <button class="modal-btn cancel" id="errorClose">OK</button>
            </div>
        </div>
    </div>
    
    <script>
        const chat=document.getElementById('chat'),input=document.getElementById('input'),sendBtn=document.getElementById('send'),dot=document.getElementById('dot'),statusEl=document.getElementById('status');
        const chatHistoryEl=document.getElementById('chatHistory');
        const deleteModal=document.getElementById('deleteModal');
        let ready=false,busy=false,currentChatId=null,chatToDelete=null;
        
        // Store stats globally so they persist
        let appStats = { sama: '-', cma: '-', total: '-' };
        
        marked.setOptions({breaks:true,gfm:true});
        
        // Close dropdowns when clicking outside
        document.addEventListener('click', (e) => {
            if(!e.target.closest('.history-item-menu') && !e.target.closest('.dropdown')) {
                document.querySelectorAll('.dropdown.show').forEach(d => d.classList.remove('show'));
            }
        });
        
        // Modal handlers
        document.getElementById('cancelDelete').addEventListener('click', () => {
            deleteModal.classList.remove('show');
            chatToDelete = null;
        });
        
        document.getElementById('confirmDelete').addEventListener('click', async () => {
            if(chatToDelete) {
                await deleteChat(chatToDelete);
                deleteModal.classList.remove('show');
                chatToDelete = null;
            }
        });
        
        async function deleteChat(chatId) {
            try {
                const r = await window.pywebview.api.delete_chat(chatId);
                if(r.success) {
                    renderChatHistory(r.chats);
                    if(chatId === currentChatId) {
                        await newChat(true);
                    }
                }
            } catch(e) {
                console.error('deleteChat error:', e);
            }
        }
        
        function showDeleteConfirm(chatId) {
            chatToDelete = chatId;
            deleteModal.classList.add('show');
        }
        
        async function init(){
            try{
                const r=await window.pywebview.api.initialize();
                if(r.status==='ready'){
                    ready=true;
                    dot.classList.remove('loading');
                    statusEl.textContent=r.total+' articles indexed';
                    // Store stats globally
                    appStats.sama = r.sama;
                    appStats.cma = r.cma;
                    appStats.total = r.total;
                    document.getElementById('s-sama').textContent=r.sama;
                    document.getElementById('s-cma').textContent=r.cma;
                    document.getElementById('s-total').textContent=r.total;
                    input.disabled=false;
                    sendBtn.disabled=false;
                    document.getElementById('overlay').classList.add('hidden');
                    input.focus();
                    // Load chat history
                    if(r.chats) renderChatHistory(r.chats);
                    // Start new chat
                    await newChat(false);
                }else{
                    statusEl.textContent='Error: '+r.message;
                }
            }catch(e){
                statusEl.textContent='Error: '+e.message;
            }
        }
        window.addEventListener('pywebviewready',init);
        
        function renderChatHistory(chats){
            let html='<div class="history-title">Recent Chats</div>';
            chats.forEach(c=>{
                const isActive=c.id===currentChatId?' active':'';
                html+=`<div class="history-item${isActive}" data-id="${c.id}">
                    <span class="history-item-text">${escHtml(c.preview)}</span>
                    <div class="history-item-menu" data-id="${c.id}">
                        <svg viewBox="0 0 24 24"><circle cx="12" cy="6" r="1.5"/><circle cx="12" cy="12" r="1.5"/><circle cx="12" cy="18" r="1.5"/></svg>
                    </div>
                    <div class="dropdown" data-id="${c.id}">
                        <div class="dropdown-item delete-chat" data-id="${c.id}">
                            <svg viewBox="0 0 24 24" stroke-width="2"><path d="M3 6h18M8 6V4a2 2 0 012-2h4a2 2 0 012 2v2m3 0v14a2 2 0 01-2 2H7a2 2 0 01-2-2V6h14"/></svg>
                            Delete
                        </div>
                    </div>
                </div>`;
            });
            chatHistoryEl.innerHTML=html;
            
            // Add click handlers for chat items
            chatHistoryEl.querySelectorAll('.history-item').forEach(el=>{
                el.addEventListener('click',(e)=>{
                    if(!e.target.closest('.history-item-menu') && !e.target.closest('.dropdown')) {
                        loadChat(el.dataset.id);
                    }
                });
            });
            
            // Add click handlers for menu buttons
            chatHistoryEl.querySelectorAll('.history-item-menu').forEach(el=>{
                el.addEventListener('click',(e)=>{
                    e.stopPropagation();
                    const dropdown = el.nextElementSibling;
                    document.querySelectorAll('.dropdown.show').forEach(d => {
                        if(d !== dropdown) d.classList.remove('show');
                    });
                    dropdown.classList.toggle('show');
                });
            });
            
            // Add click handlers for delete buttons
            chatHistoryEl.querySelectorAll('.delete-chat').forEach(el=>{
                el.addEventListener('click',(e)=>{
                    e.stopPropagation();
                    showDeleteConfirm(el.dataset.id);
                });
            });
        }
        
        async function newChat(updateUI=true){
            try{
                const r=await window.pywebview.api.new_chat();
                currentChatId=r.id;
                if(updateUI){
                    // Restore full welcome screen with stats from global appStats
                    chat.innerHTML='<div class="welcome" id="welcome"><div class="welcome-icon"><svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M12 2L2 7l10 5 10-5-10-5zM2 17l10 5 10-5M2 12l10 5 10-5" stroke="currentColor" fill="none"/></svg></div><h1>TadqeeqAI</h1><p>Bilingual AI assistant for Saudi Arabian financial regulations. Query SAMA and CMA documents in English or Arabic with accurate, cited responses.</p><div class="welcome-stats"><div class="stat"><div class="stat-val" id="s-sama">' + appStats.sama + '</div><div class="stat-lbl">SAMA</div></div><div class="stat"><div class="stat-val" id="s-cma">' + appStats.cma + '</div><div class="stat-lbl">CMA</div></div><div class="stat"><div class="stat-val" id="s-total">' + appStats.total + '</div><div class="stat-lbl">Articles</div></div></div></div>';
                    renderChatHistory(r.chats);
                    // Clear document badge
                    const docBadge = document.getElementById('docBadge');
                    if (docBadge) { docBadge.style.display = 'none'; docBadge.innerHTML = ''; }
                }
            }catch(e){
                console.error('newChat error:',e);
            }
        }
        
        async function loadChat(chatId){
            try{
                const r=await window.pywebview.api.load_chat(chatId);
                currentChatId=chatId;
                // Clear and render messages
                chat.innerHTML='';
                r.messages.forEach(m=>{
                    addMsg(m.role,m.content,m.sources||null,m.regulator||null);
                });
                // Update active state
                chatHistoryEl.querySelectorAll('.history-item').forEach(el=>{
                    el.classList.toggle('active',el.dataset.id===chatId);
                });
            }catch(e){
                console.error('loadChat error:',e);
            }
        }
        
        document.getElementById('newChatBtn').addEventListener('click',()=>newChat(true));
        
        input.addEventListener('input',()=>{input.style.height='auto';input.style.height=Math.min(input.scrollHeight,120)+'px';});
        input.addEventListener('keydown',e=>{if(e.key==='Enter'&&!e.shiftKey){e.preventDefault();send();}});
        sendBtn.addEventListener('click',send);
        document.querySelectorAll('.ex').forEach(el=>{el.addEventListener('click',()=>{input.value=el.dataset.q;send();});});
        
        function escHtml(t){const d=document.createElement('div');d.textContent=t;return d.innerHTML;}
        function renderMd(text){try{return marked.parse(text);}catch(e){return escHtml(text);}}
        function isArabic(text){return /[\u0600-\u06FF]/.test(text)&&(text.match(/[\u0600-\u06FF]/g)||[]).length>text.length*0.3;}
        
        const logoSvg='<svg viewBox="0 0 24 24" fill="none"><path d="M12 2L2 7l10 5 10-5-10-5zM2 17l10 5 10-5M2 12l10 5 10-5" stroke="currentColor" stroke-width="2" fill="none"/></svg>';
        
        function addMsg(role,text,sources,reg){
            const w=document.getElementById('welcome');if(w)w.style.display='none';
            const div=document.createElement('div');
            div.className='msg '+role;
            let regHtml=reg&&reg!=='NONE'?'<span class="reg '+reg.toLowerCase()+'">'+reg+'</span>':'';
            let srcHtml='';
            if(sources&&sources.length){
                srcHtml='<div class="sources"><div class="src-title">Sources</div>';
                sources.forEach(s=>{srcHtml+='<span class="src">'+escHtml(s.article)+'</span>';});
                srcHtml+='</div>';
            }
            const textHtml=role==='assistant'?renderMd(text):escHtml(text);
            const rtlClass=isArabic(text)?' rtl':'';
            const avatarContent=role==='user'?'You':logoSvg;
            div.innerHTML='<div class="avatar">'+avatarContent+'</div><div class="msg-body"><div class="msg-header"><span class="msg-role">'+(role==='user'?'You':'TadqeeqAI')+'</span>'+regHtml+'</div><div class="msg-text'+rtlClass+'">'+textHtml+'</div>'+srcHtml+'</div>';
            chat.appendChild(div);
            chat.scrollTop=chat.scrollHeight;
        }
        
        function addLoading(){
            const w=document.getElementById('welcome');if(w)w.style.display='none';
            const div=document.createElement('div');
            div.className='msg assistant';
            div.id='loading';
            div.innerHTML='<div class="avatar">'+logoSvg+'</div><div class="msg-body"><div class="msg-header"><span class="msg-role">TadqeeqAI</span></div><div class="loading-msg"><div class="dots"><span></span><span></span><span></span></div><span>Searching regulations...</span></div></div>';
            chat.appendChild(div);
            chat.scrollTop=chat.scrollHeight;
        }
        
        async function send(){
            const q=input.value.trim();
            if(!q||busy||!ready)return;
            busy=true;
            sendBtn.disabled=true;
            input.disabled=true;
            input.value='';
            input.style.height='auto';
            addMsg('user',q,null,null);
            addLoading();
            try{
                const r=await window.pywebview.api.query(q);
                document.getElementById('loading')?.remove();
                addMsg('assistant',r.answer,r.sources,r.regulator);
                // Refresh chat history
                const chats=await window.pywebview.api.get_chats();
                renderChatHistory(chats.chats);
            }catch(e){
                document.getElementById('loading')?.remove();
                addMsg('assistant','An error occurred while processing your request.',null,null);
            }
            busy=false;
            sendBtn.disabled=false;
            input.disabled=false;
            input.focus();
        }
        
        // Error modal helper
        function showError(title, message) {
            const modal = document.getElementById('errorModal');
            document.getElementById('errorTitle').textContent = title;
            document.getElementById('errorText').textContent = message;
            modal.classList.add('show');
        }
        document.getElementById('errorClose').addEventListener('click', () => {
            document.getElementById('errorModal').classList.remove('show');
        });
        
        // ========== V2.2 FEATURES ==========
        (function initV22() {
            // Get elements
            const exportBtn = document.getElementById('exportBtn');
            const exportMenu = document.getElementById('exportMenu');
            const exportMd = document.getElementById('exportMd');
            const exportPdf = document.getElementById('exportPdf');
            const dropOverlay = document.getElementById('dropOverlay');
            const fileInput = document.getElementById('fileInput');
            const attachBtn = document.getElementById('attachBtn');
            const docBadge = document.getElementById('docBadge');
            const mainEl = document.querySelector('.main');
            
            // Export menu toggle
            if (exportBtn && exportMenu) {
                exportBtn.addEventListener('click', (e) => {
                    e.stopPropagation();
                    exportMenu.classList.toggle('show');
                });
                document.addEventListener('click', (e) => {
                    if (!e.target.closest('#exportBtn') && !e.target.closest('#exportMenu')) {
                        exportMenu.classList.remove('show');
                    }
                });
            }
            
            // Export Markdown
            if (exportMd) {
                exportMd.addEventListener('click', async () => {
                    if (exportMenu) exportMenu.classList.remove('show');
                    try {
                        console.log('Exporting markdown...');
                        const r = await window.pywebview.api.export_markdown();
                        console.log('Export result:', r);
                        if (r.error) { 
                            showError('Export Failed', r.error); 
                            return; 
                        }
                        if (r.success) {
                            // File was saved successfully
                            showError('Export Complete', 'Chat exported to: ' + r.path);
                        }
                    } catch (e) { 
                        console.error('Export error:', e);
                        showError('Export Failed', e.message || 'Unknown error'); 
                    }
                });
            }
            
            // Export PDF
            if (exportPdf) {
                exportPdf.addEventListener('click', async () => {
                    if (exportMenu) exportMenu.classList.remove('show');
                    try {
                        console.log('Exporting PDF...');
                        const r = await window.pywebview.api.export_pdf();
                        console.log('Export result:', r);
                        if (r.error) { 
                            showError('Export Failed', r.error); 
                            return; 
                        }
                        if (r.success) {
                            // File was saved successfully
                            showError('Export Complete', 'Chat exported to: ' + r.path);
                        }
                    } catch (e) { 
                        console.error('Export error:', e);
                        showError('Export Failed', e.message || 'Unknown error'); 
                    }
                });
            }
            
            // Drag and drop
            if (mainEl && dropOverlay) {
                mainEl.addEventListener('dragenter', (e) => { e.preventDefault(); dropOverlay.classList.add('active'); });
                mainEl.addEventListener('dragover', (e) => { e.preventDefault(); dropOverlay.classList.add('active'); });
                dropOverlay.addEventListener('dragleave', (e) => { e.preventDefault(); dropOverlay.classList.remove('active'); });
                dropOverlay.addEventListener('drop', (e) => {
                    e.preventDefault();
                    dropOverlay.classList.remove('active');
                    if (e.dataTransfer.files[0]) handleFile(e.dataTransfer.files[0]);
                });
            }
            
            // Attach button
            if (attachBtn && fileInput) {
                attachBtn.addEventListener('click', () => fileInput.click());
                fileInput.addEventListener('change', (e) => {
                    if (e.target.files[0]) handleFile(e.target.files[0]);
                });
            }
            
            // Helper functions
            function downloadFile(content, filename, type) {
                const blob = new Blob([content], { type });
                downloadBlob(blob, filename);
            }
            
            function downloadBlob(blob, filename) {
                const url = URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                a.download = filename;
                document.body.appendChild(a);
                a.click();
                document.body.removeChild(a);
                URL.revokeObjectURL(url);
            }
            
            function b64toBlob(b64Data, contentType) {
                const byteChars = atob(b64Data);
                const byteArrays = [];
                for (let offset = 0; offset < byteChars.length; offset += 512) {
                    const slice = byteChars.slice(offset, offset + 512);
                    const byteNumbers = new Array(slice.length);
                    for (let i = 0; i < slice.length; i++) byteNumbers[i] = slice.charCodeAt(i);
                    byteArrays.push(new Uint8Array(byteNumbers));
                }
                return new Blob(byteArrays, { type: contentType });
            }
            
            async function handleFile(file) {
                const ext = file.name.split('.').pop().toLowerCase();
                if (!['pdf', 'docx', 'doc'].includes(ext)) {
                    showError('Invalid File', 'Please upload a PDF or DOCX file.');
                    return;
                }
                const reader = new FileReader();
                reader.onload = async (e) => {
                    try {
                        const base64 = e.target.result.split(',')[1];
                        const r = await window.pywebview.api.upload_document(base64, file.name);
                        if (r.error) { showError('Upload Failed', r.error); return; }
                        showDocBadge(r);
                        const w = document.getElementById('welcome');
                        if (w) w.style.display = 'none';
                        addMsg('assistant', '📄 **Document loaded:** ' + r.filename + '\\n\\n**Pages:** ' + r.pages + ' | **Characters:** ' + r.chars.toLocaleString() + '\\n**Type:** ' + r.summary.type + (r.summary.has_arabic ? '\\n**Language:** Contains Arabic text' : '') + '\\n\\nYou can now ask questions about this document or click **Check Compliance** to run an automated scan.', null, null);
                    } catch (err) { showError('Upload Failed', err.message); }
                };
                reader.readAsDataURL(file);
            }
            
            function showDocBadge(info) {
                if (!docBadge) return;
                docBadge.innerHTML = '<div class="doc-badge"><svg viewBox="0 0 24 24" stroke-width="2"><path d="M14 2H6a2 2 0 00-2 2v16a2 2 0 002 2h12a2 2 0 002-2V8z"/><polyline points="14 2 14 8 20 8"/></svg><div class="doc-badge-info"><div class="doc-badge-name">' + escHtml(info.filename) + '</div><div class="doc-badge-meta">' + info.pages + ' pages \\u00b7 ' + info.summary.type + '</div></div><div class="doc-badge-actions"><button class="doc-badge-btn check" onclick="window.runComplianceCheck()">Check Compliance</button><button class="doc-badge-btn clear" onclick="window.clearDocument()">\\u2715</button></div></div>';
                docBadge.style.display = 'block';
            }
            
            // Global functions for onclick handlers
            window.clearDocument = async function() {
                try {
                    await window.pywebview.api.clear_document();
                    if (docBadge) { docBadge.style.display = 'none'; docBadge.innerHTML = ''; }
                } catch (e) { console.error('Clear failed:', e); }
            };
            
            window.runComplianceCheck = async function() {
                if (busy) return;
                busy = true;
                addMsg('assistant', '🔍 Running compliance check...', null, null);
                try {
                    const r = await window.pywebview.api.run_compliance_check();
                    const loading = document.getElementById('loading');
                    if (loading) loading.remove();
                    const msgs = chat.querySelectorAll('.msg.assistant');
                    if (msgs.length) msgs[msgs.length - 1].remove();
                    if (r.error) { addMsg('assistant', '❌ ' + r.error, null, null); }
                    else { showComplianceReport(r); }
                } catch (e) { addMsg('assistant', '❌ Error: ' + e.message, null, null); }
                busy = false;
            };
            
            function showComplianceReport(r) {
                const scoreClass = r.score >= 80 ? 'good' : r.score >= 50 ? 'warning' : 'bad';
                let itemsHtml = '';
                r.checks.forEach(c => {
                    const iconClass = c.status === 'compliant' ? 'compliant' : 'warning';
                    const icon = c.status === 'compliant'
                        ? '<svg viewBox="0 0 24 24" stroke-width="2"><polyline points="20 6 9 17 4 12"/></svg>'
                        : '<svg viewBox="0 0 24 24" stroke-width="2"><circle cx="12" cy="12" r="10"/><line x1="12" y1="8" x2="12" y2="12"/><line x1="12" y1="16" x2="12.01" y2="16"/></svg>';
                    itemsHtml += '<div class="compliance-item"><div class="compliance-icon ' + iconClass + '">' + icon + '</div><div class="compliance-content"><div class="compliance-name">' + c.name + '</div><div class="compliance-reg">' + c.regulation + '</div><div class="compliance-detail">' + c.detail + '</div></div></div>';
                });
                const reportHtml = '<div class="compliance-report"><div class="compliance-header"><div class="compliance-title">Compliance Report: ' + escHtml(r.filename) + '</div><div class="compliance-score ' + scoreClass + '">' + r.score + '% Compliant</div></div>' + itemsHtml + '</div>';
                const div = document.createElement('div');
                div.innerHTML = reportHtml;
                chat.appendChild(div);
                chat.scrollTop = chat.scrollHeight;
            }
        })();
    </script>
</body>
</html>'''


if __name__ == '__main__':
    print("\n" + "="*50)
    print("   TADQEEQAI v2.2")
    print("   Hybrid Search: BM25 + Semantic")
    print("   SAMA + CMA · English + Arabic")
    print("   Document Analysis · Export")
    print("="*50 + "\n")
    api = API()
    window = webview.create_window('TadqeeqAI v2.2', html=HTML, js_api=api, width=1200, height=800, min_size=(900,600), background_color='#0f1419', text_select=True)
    api.set_window(window)
    webview.start(debug=False)
